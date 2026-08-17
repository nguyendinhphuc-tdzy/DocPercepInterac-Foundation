# anonymize_docx.py
# Dùng: python anonymize_docx.py input.docx output_anonymized.docx
#
# Chạy CỤC BỘ trên máy bạn — không gửi file đi đâu cả.
# Chỉ dùng python-docx (đã duyệt CRADL) — không có network call nào trong script này.

import sys
from docx import Document
from anonymize_common import anonymize_text


def _process_paragraph(paragraph):
    """Thay text ở TỪNG RUN, không thay cả paragraph.text — để giữ nguyên
    định dạng (bold/italic/font) của từng đoạn chữ, không làm mất style."""
    for run in paragraph.runs:
        if run.text:
            run.text = anonymize_text(run.text)


def _process_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _process_paragraph(paragraph)
            # Bảng lồng bảng (nested table) — xử lý đệ quy
            for nested_table in cell.tables:
                _process_table(nested_table)


def anonymize_docx(input_path: str, output_path: str):
    doc = Document(input_path)

    # 1. Toàn bộ đoạn văn trong thân tài liệu
    for paragraph in doc.paragraphs:
        _process_paragraph(paragraph)

    # 2. Toàn bộ bảng trong thân tài liệu
    for table in doc.tables:
        _process_table(table)

    # 3. Header & Footer — DỄ BỊ BỎ SÓT nếu chỉ xử lý body, nhưng lại hay chứa
    # tên client (vd: "Prepared for [Client Name]" ở header mọi trang)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _process_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            _process_paragraph(paragraph)
        for table in section.header.tables:
            _process_table(table)
        for table in section.footer.tables:
            _process_table(table)

    doc.save(output_path)
    print(f"Đã ẩn danh: {input_path} -> {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Dùng: python anonymize_docx.py input.docx output_anonymized.docx")
        sys.exit(1)
    anonymize_docx(sys.argv[1], sys.argv[2])
