"""P3-04 duplicate_ordinal stress test — synthetic, non-financial, always
runs (no `skipif`).

tests/test_anchor_builder.py::test_p304_docx_duplicate_ordinal_survives_uneven_drift_on_real_document
covers the same failure mode, but only against a real client DOCX gated by
`@requires_real_docx` — on a clean checkout or in CI without that file, it
is silently skipped, leaving zero coverage for the project's most critical
anchor-stability guarantee (Foundation_Build_Plan_v3.md §9.2's mandatory
milestone). fixture_anchor_stress.docx
(tests/fixtures/_generate_anchor_stress_docx.py) reproduces the same
structural ambiguity — a caption repeated many times with an identical
(style_id, text) signature, with other content between each repetition —
entirely synthetically and non-financially, so this test runs
unconditionally. The real-doc test above is kept as-is (not modified,
not deleted) as a supplementary real-world check; this file does not
replace it.

Does NOT modify perception/anchor_builder.py — duplicate_ordinal already
exists and is already correct there. This only adds fixture + coverage.
"""
import importlib.util
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document as DocxDocument  # noqa: E402

from perception.anchor_builder import assign_anchors, resolve_docx_anchor  # noqa: E402
from perception.parser import parse_docx  # noqa: E402

FIXTURES = Path(__file__).resolve().parent / "fixtures"
FIXTURE = FIXTURES / "fixture_anchor_stress.docx"

# Load the generator module by path (not via package import) so the
# fixture's structural constants live in exactly one place and these tests
# can't drift from what tests/fixtures/_generate_anchor_stress_docx.py
# actually produced — same pattern as test_parser_generic.py.
_spec = importlib.util.spec_from_file_location(
    "_generate_anchor_stress_docx", FIXTURES / "_generate_anchor_stress_docx.py"
)
_gen = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(_gen)
CAPTION_TEXT = _gen.CAPTION_TEXT
CAPTION_STYLE = _gen.CAPTION_STYLE
CAPTION_COUNT = _gen.CAPTION_COUNT


def _insert_paragraph_before(doc, index: int, text: str):
    """Insert a new paragraph immediately before doc.paragraphs[index]."""
    anchor_para = doc.paragraphs[index]
    new_para = doc.add_paragraph(text)  # appends at the end
    new_para._p.getparent().remove(new_para._p)
    anchor_para._p.addprevious(new_para._p)


def test_fixture_exists():
    assert FIXTURE.exists(), (
        "fixture_anchor_stress.docx is missing — regenerate it with "
        "python tests/fixtures/_generate_anchor_stress_docx.py"
    )


def test_synthetic_fixture_has_duplicate_caption_signature_and_correct_ordinals():
    """Guards the fixture itself (>= 6 identical-signature occurrences —
    the ambiguity duplicate_ordinal exists for) and that assign_anchors
    ranks each occurrence correctly, including a middle one (ordinal 4)."""
    blocks = parse_docx(str(FIXTURE))

    occurrences = [
        b for b in blocks if b["text"] == CAPTION_TEXT and b["style_id"] == CAPTION_STYLE
    ]
    assert len(occurrences) == CAPTION_COUNT
    assert len(occurrences) >= 6, "fixture must have >= 6 identical-signature occurrences"

    anchors = assign_anchors(blocks, "docx")
    for k in (0, 4, CAPTION_COUNT - 1):
        block = occurrences[k]
        anchor = anchors[blocks.index(block)]
        assert anchor.duplicate_ordinal == k


def test_p304_synthetic_duplicate_ordinal_survives_uneven_drift_between_occurrences(tmp_path):
    """CORE P3-04, synthetic: insert paragraphs *between* two occurrences of
    the duplicated caption (not at the very start) — this shifts only the
    target occurrence (and everything after it), while the *previous*
    occurrence stays completely unmoved and becomes numerically closer to
    the anchor's originally recorded paragraph_index. A tie-break based on
    "nearest paragraph_index" picks that unmoved neighbor — the wrong
    occurrence. duplicate_ordinal (rank among same-signature paragraphs,
    immune to uneven shifts) exists to fix exactly this.
    """
    blocks = parse_docx(str(FIXTURE))
    occurrences = [
        b for b in blocks if b["text"] == CAPTION_TEXT and b["style_id"] == CAPTION_STYLE
    ]
    assert len(occurrences) >= 6

    target_ordinal = 4  # a middle occurrence, not the first/last
    target_block = occurrences[target_ordinal]
    target_index = target_block["paragraph_index"]
    previous_occurrence_index = occurrences[target_ordinal - 1]["paragraph_index"]

    anchors = assign_anchors(blocks, "docx")
    target_anchor = anchors[blocks.index(target_block)]
    assert target_anchor.duplicate_ordinal == target_ordinal

    # Insert filler paragraphs immediately before the target occurrence —
    # shifts the target (and everything after) forward by INSERT_COUNT,
    # while the previous occurrence of the same caption is left untouched.
    INSERT_COUNT = 50
    doc = DocxDocument(str(FIXTURE))
    for i in range(INSERT_COUNT):
        _insert_paragraph_before(doc, target_index, f"Injected filler paragraph {i}.")
    drifted_path = tmp_path / "drifted.docx"
    doc.save(drifted_path)

    doc_for_resolve = DocxDocument(str(drifted_path))
    expected_target = doc_for_resolve.paragraphs[target_index + INSERT_COUNT]
    assert expected_target.text == CAPTION_TEXT

    # Sanity / false-pass guard: a naive "nearest to the originally
    # recorded paragraph_index" tie-break must pick the WRONG occurrence
    # here (the unmoved previous one) — proving this scenario isn't
    # trivially solved and that the assertion below is actually exercising
    # duplicate_ordinal's disambiguation, not passing by coincidence.
    same_signature_paragraphs = [
        (i, p)
        for i, p in enumerate(doc_for_resolve.paragraphs)
        if p.text == CAPTION_TEXT and (p.style.style_id if p.style else "") == CAPTION_STYLE
    ]
    naive_pick_index, _naive_pick = min(
        same_signature_paragraphs, key=lambda pair: abs(pair[0] - target_index)
    )
    assert naive_pick_index == previous_occurrence_index, (
        "test setup no longer reproduces the failure mode duplicate_ordinal "
        "fixes — a naive nearest-paragraph_index tie-break should pick the "
        "unmoved previous occurrence, not the actual (shifted) target"
    )
    assert naive_pick_index != target_index + INSERT_COUNT

    # The real fix: resolving via duplicate_ordinal must find the ACTUAL
    # target, not the nearer neighbor the naive strategy above picked.
    resolved, message = resolve_docx_anchor(doc_for_resolve, target_anchor)

    assert resolved.text == CAPTION_TEXT
    assert resolved._p is expected_target._p, (
        "resolved the wrong occurrence of the duplicated caption — likely "
        "picked the nearer, unrelated neighbor instead of the actual target"
    )
    assert message is None  # Strategy 1, ordinal-disambiguated — no warning
