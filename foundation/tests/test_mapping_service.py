"""Unit tests for applications/gpts/mapping_service.py — the orchestration
layer behind POST /api/process.
"""
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from applications.gpts.mapping_service import run_mapping  # noqa: E402

DEMO_ROOT = Path(__file__).resolve().parents[2] / "anonymize client" / "Demo files" / "Demo files"
SOURCE_XLSX = DEMO_ROOT / "FA&RPTS & Appendix I" / "FA&RPTs" / "HMV-FA&RPT FY2024.xlsx"
TARGET_DOCX = DEMO_ROOT / "Compare LF" / "HMV-26-Final-Local File for FY2024-EN-R2901KPMG_drifted.docx"

requires_demo_fixtures = pytest.mark.skipif(
    not (SOURCE_XLSX.exists() and TARGET_DOCX.exists()),
    reason="HMV demo fixture files not present on this machine",
)


@requires_demo_fixtures
def test_run_mapping_maps_all_demo_rules(tmp_path):
    result = run_mapping([str(SOURCE_XLSX)], str(TARGET_DOCX), str(tmp_path))

    assert len(result.source_elements) > 0
    assert len(result.target_elements) > 0
    assert len(result.mapped) == 3

    mapped_source_anchors = {m.source_anchor for m in result.mapped}
    assert mapped_source_anchors == {"RPTs!E8", "RPTs!F8", "Financial Analysis!D7"}

    assert result.patched_docx_path is not None
    assert Path(result.patched_docx_path).exists()

    # Every mapped entry must resolve to a real element in target_elements,
    # by table_hash (self-heal-aware) — not a stale table_index. This
    # fixture's DEMO_RULES targets genuinely self-heal (table 6 -> 4 in
    # the drifted file), so this also guards against regressing to naive
    # table_index matching (frontend/src/components/results/ResultsPane.tsx
    # relies on target_element_index instead of re-deriving this). The
    # resolved element's own text is still empty at this point — it's the
    # pristine pre-patch parse, DEMO_RULES fills these cells in for the
    # *output file*, not for this in-memory Element — so only the anchor
    # identity is checked, not its text.
    target_by_index = {e.index: e for e in result.target_elements}
    for m in result.mapped:
        assert m.target_element_index is not None, f"no element resolved for {m.target_anchor}"
        assert m.target_element_index in target_by_index
        resolved_anchor = target_by_index[m.target_element_index].anchor
        assert resolved_anchor.format == "docx" and resolved_anchor.table_index is not None


def test_run_mapping_returns_zero_matches_for_unrelated_pair(tmp_path):
    import openpyxl
    from docx import Document as DocxDocument

    xlsx_path = tmp_path / "unrelated.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "Not part of DEMO_RULES"
    wb.save(xlsx_path)
    wb.close()

    docx_path = tmp_path / "unrelated.docx"
    doc = DocxDocument()
    doc.add_paragraph("Just a paragraph, no tables.")
    doc.save(docx_path)

    result = run_mapping([str(xlsx_path)], str(docx_path), str(tmp_path / "out"))

    assert len(result.source_elements) == 1
    assert len(result.target_elements) == 1
    assert result.mapped == []
    assert result.patched_docx_path is None


def test_run_mapping_merges_elements_from_multiple_source_files(tmp_path):
    """H3 fix: uploading more than one source file must not silently drop
    everything after the first — a real use case for GTPS's own demo
    (FA&RPTs + Appendix I together), and for any other function that needs
    more than one input document."""
    import openpyxl

    xlsx_path_1 = tmp_path / "first.xlsx"
    wb1 = openpyxl.Workbook()
    wb1.active["A1"] = "From first file"
    wb1.save(xlsx_path_1)
    wb1.close()

    xlsx_path_2 = tmp_path / "second.xlsx"
    wb2 = openpyxl.Workbook()
    wb2.active["A1"] = "From second file"
    wb2.save(xlsx_path_2)
    wb2.close()

    from docx import Document as DocxDocument

    docx_path = tmp_path / "target.docx"
    DocxDocument().save(docx_path)

    result = run_mapping([str(xlsx_path_1), str(xlsx_path_2)], str(docx_path), str(tmp_path / "out"))

    texts = {e.text for e in result.source_elements}
    assert texts == {"From first file", "From second file"}
    indices = [e.index for e in result.source_elements]
    assert indices == sorted(set(indices)) and len(indices) == len(set(indices))  # unique, no collision
