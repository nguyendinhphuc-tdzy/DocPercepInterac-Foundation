"""Tests for the Comprehensive Document Perception & Media Layer phase:
images/charts/drawings/headers/footers/footnotes/endnotes/comments/
annotations as first-class Elements, the media manifest, and
capability metadata (`detected`/`extracted`/`rendered`/`selectable`/
`editable`). These assert specific type counts and relationships, not
just `len(elements) > 0` — see perception/parser.py, anchor_builder.py,
element_classifier.py, models.py.
"""
import io
import sys
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from perception.anchor_builder import assign_anchors
from perception.element_classifier import classify_blocks
from perception.models import ElementType, ExtractionLevel
from perception.parser import (
    extract_geometry,
    extract_media_manifest,
    extract_worksheet_metadata,
    resolve_media_bytes,
)

FIXTURES = Path(__file__).resolve().parent / "fixtures"


def _perceive(path: str, fmt: str):
    blocks = extract_geometry(path)
    anchors = assign_anchors(blocks, fmt)
    return blocks, classify_blocks(blocks, fmt, anchors)


# --- DOCX images/drawings (real fixture) ------------------------------------


def test_docx_images_are_first_class_elements_with_media_manifest_entries():
    path = str(FIXTURES / "fixture_bcdt.docx")
    _, elements = _perceive(path, "docx")

    images = [e for e in elements if e.type == ElementType.IMAGE]
    assert len(images) >= 1, "fixture_bcdt.docx is known to embed images (signature graphics)"
    for img in images:
        assert img.capabilities.detected is True
        assert img.capabilities.extracted == ExtractionLevel.FULL
        assert img.capabilities.editable is False  # images are inspect/select only, never fake-editable
        assert img.anchor.format == "docx"
        assert img.anchor.relationship_id is not None
        assert img.anchor.media_id == img.anchor.relationship_id  # docx media_id IS the relationship id

    media = extract_media_manifest(path, "docx")
    media_ids = {m.media_id for m in media}
    assert media_ids, "media manifest should be non-empty for a document with embedded images"
    for img in images:
        assert img.anchor.media_id in media_ids, "every image Element's media_id must resolve in the manifest"


def test_docx_media_bytes_resolve_and_are_non_empty():
    path = str(FIXTURES / "fixture_bcdt.docx")
    media = extract_media_manifest(path, "docx")
    assert media
    resolved = resolve_media_bytes(path, "docx", media[0].media_id)
    assert resolved is not None
    data, mime_type = resolved
    assert len(data) > 0
    assert mime_type.startswith("image/")


def test_docx_unresolvable_media_id_returns_none_not_a_guess():
    path = str(FIXTURES / "fixture_bcdt.docx")
    assert resolve_media_bytes(path, "docx", "not-a-real-relationship-id") is None


def test_docx_images_preserve_reading_order_interleaved_with_paragraphs():
    """An image block's paragraph_index must place it among the paragraphs
    it actually sits between — never bucketed at the end of the list
    regardless of true position (the bug this phase's reading-order
    requirement explicitly targets)."""
    path = str(FIXTURES / "fixture_bcdt.docx")
    blocks, _ = _perceive(path, "docx")
    kinds_in_order = [b["kind"] for b in blocks]
    image_positions = [i for i, k in enumerate(kinds_in_order) if k == "image"]
    assert image_positions
    # At least one image must NOT be the very last block — i.e. it isn't
    # simply appended after everything else the way the pre-phase
    # "collect paragraphs, then collect media" approach would.
    assert any(pos != len(blocks) - 1 for pos in image_positions) or len(blocks) == 1


# --- DOCX headers/footers/footnotes/endnotes/comments (synthetic fixture) --


def _build_docx_with_chrome(tmp_path) -> Path:
    """A real docx (python-docx: paragraph + header + footer) with
    footnotes.xml/endnotes.xml/comments.xml parts injected afterward —
    python-docx has no API for authoring those three, so this test
    exercises the same raw-OOXML-part reading path
    (_docx_notes_and_comments) against hand-built parts, which is exactly
    what that function must survive regardless of which tool authored the
    original document (Word, LibreOffice, ...)."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Body paragraph.")
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Header text"
    section.footer.paragraphs[0].text = "Footer text"
    path = tmp_path / "chrome.docx"
    doc.save(path)

    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    footnotes_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="{w_ns}">
  <w:footnote w:type="separator" w:id="0"><w:p/></w:footnote>
  <w:footnote w:id="1"><w:p><w:r><w:t>A real footnote.</w:t></w:r></w:p></w:footnote>
</w:footnotes>'''
    endnotes_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:endnotes xmlns:w="{w_ns}">
  <w:endnote w:type="separator" w:id="0"><w:p/></w:endnote>
  <w:endnote w:id="1"><w:p><w:r><w:t>A real endnote.</w:t></w:r></w:p></w:endnote>
</w:endnotes>'''
    comments_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="{w_ns}">
  <w:comment w:id="1" w:author="Reviewer"><w:p><w:r><w:t>A real comment.</w:t></w:r></w:p></w:comment>
</w:comments>'''

    # Rewrite the zip with the three extra parts added — python's zipfile
    # can't append to an existing archive member-by-member in place, so
    # copy every existing entry across and add the new ones.
    buf = io.BytesIO()
    with zipfile.ZipFile(path, "r") as src, zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            dst.writestr(item, src.read(item.filename))
        dst.writestr("word/footnotes.xml", footnotes_xml)
        dst.writestr("word/endnotes.xml", endnotes_xml)
        dst.writestr("word/comments.xml", comments_xml)
    path.write_bytes(buf.getvalue())
    return path


def test_docx_chrome_objects_are_detected_and_classified(tmp_path):
    path = _build_docx_with_chrome(tmp_path)
    _, elements = _perceive(str(path), "docx")
    by_type = {}
    for e in elements:
        by_type.setdefault(e.type, []).append(e)

    assert ElementType.HEADER in by_type and by_type[ElementType.HEADER][0].text == "Header text"
    assert ElementType.FOOTER in by_type and by_type[ElementType.FOOTER][0].text == "Footer text"
    assert ElementType.FOOTNOTE in by_type and by_type[ElementType.FOOTNOTE][0].text == "A real footnote."
    assert ElementType.ENDNOTE in by_type and by_type[ElementType.ENDNOTE][0].text == "A real endnote."
    assert ElementType.COMMENT in by_type and by_type[ElementType.COMMENT][0].text == "A real comment."

    # None of these are silently dropped, but none are falsely claimed
    # editable or selectable either — honest partial support.
    for etype in (ElementType.HEADER, ElementType.FOOTER, ElementType.FOOTNOTE, ElementType.ENDNOTE, ElementType.COMMENT):
        el = by_type[etype][0]
        assert el.capabilities.detected is True
        assert el.capabilities.editable is False
        assert el.capabilities.selectable is False

    # The separator/continuation placeholder (w:id="0") must be excluded.
    footnote_ids = {e.anchor.drawing_id for e in by_type[ElementType.FOOTNOTE]}
    assert "0" not in footnote_ids


# --- XLSX images/charts/merged/hidden/freeze (synthetic fixture) -----------


def _build_xlsx_with_drawings(tmp_path) -> Path:
    import openpyxl
    from openpyxl.chart import BarChart, Reference
    from openpyxl.drawing.image import Image
    from PIL import Image as PILImage

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    for r in range(1, 6):
        ws.cell(row=r, column=1, value=f"Label {r}")
        ws.cell(row=r, column=2, value=r * 10)

    ws.merge_cells("A1:B1")
    ws.row_dimensions[3].hidden = True
    ws.column_dimensions["C"].hidden = True
    ws.freeze_panes = "A2"

    chart = BarChart()
    chart.title = "Test Chart"
    data = Reference(ws, min_col=2, min_row=1, max_row=5)
    chart.add_data(data)
    ws.add_chart(chart, "E2")

    # A minimal real PNG (openpyxl.drawing.image.Image needs a real decodable image).
    img_buf = io.BytesIO()
    PILImage.new("RGB", (40, 20), color=(200, 50, 50)).save(img_buf, format="PNG")
    img_buf.seek(0)
    ws.add_image(Image(img_buf), "E10")

    path = tmp_path / "drawings.xlsx"
    wb.save(path)
    wb.close()
    return path


def test_xlsx_images_and_charts_are_first_class_elements(tmp_path):
    path = str(_build_xlsx_with_drawings(tmp_path))
    _, elements = _perceive(path, "xlsx")

    images = [e for e in elements if e.type == ElementType.IMAGE]
    charts = [e for e in elements if e.type == ElementType.CHART]
    assert len(images) == 1
    assert len(charts) == 1

    img = images[0]
    assert img.capabilities.editable is False
    assert img.anchor.format == "xlsx"
    assert img.anchor.from_cell is not None
    assert img.anchor.media_id is not None

    chart = charts[0]
    assert chart.name == "Test Chart" or "Chart" in chart.name
    assert chart.capabilities.extracted == ExtractionLevel.PARTIAL  # series data not deeply parsed this phase
    assert chart.anchor.media_id is None  # charts have no servable raster — no manifest entry

    media = extract_media_manifest(path, "xlsx")
    assert len(media) == 1
    assert media[0].media_id == img.anchor.media_id

    resolved = resolve_media_bytes(path, "xlsx", media[0].media_id)
    assert resolved is not None
    data, mime_type = resolved
    assert len(data) > 0
    assert mime_type == "image/png"


def test_xlsx_worksheet_metadata_reports_merges_hidden_and_freeze(tmp_path):
    path = str(_build_xlsx_with_drawings(tmp_path))
    metadata = extract_worksheet_metadata(path)
    assert len(metadata) == 1
    ws_meta = metadata[0]
    assert "A1:B1" in ws_meta.merged_ranges
    assert 3 in ws_meta.hidden_rows
    assert "C" in ws_meta.hidden_columns
    assert ws_meta.freeze_panes == "A2"


def test_xlsx_formula_is_preserved_alongside_displayed_value(tmp_path):
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = 10
    ws["A2"] = 20
    ws["A3"] = "=A1+A2"
    path = tmp_path / "formula.xlsx"
    wb.save(path)
    wb.close()

    blocks, elements = _perceive(str(path), "xlsx")
    formula_block = next(b for b in blocks if b["cell_address"] == "A3")
    assert (formula_block.get("extra") or {}).get("formula") == "=A1+A2"
    # openpyxl only caches a computed value if the file was last saved by a
    # real spreadsheet app (Excel/LibreOffice) — a file saved by openpyxl
    # itself has no cached result, so the displayed text legitimately falls
    # back to the formula string itself here. What matters for this test is
    # that the formula survived at all (not silently dropped in favor of
    # only the cached value, which was the pre-phase `data_only=True`-only
    # behavior).
    formula_element = next(e for e in elements if e.anchor.format == "xlsx" and e.anchor.cell_address == "A3")
    assert formula_element.text  # non-empty either way


# --- PDF images/annotations (real fixtures) ---------------------------------


def test_pdf_annotations_detected_from_real_fixture():
    path = str(FIXTURES / "fixture_report_2.pdf")
    _, elements = _perceive(path, "pdf")
    annotations = [e for e in elements if e.type == ElementType.ANNOTATION]
    assert len(annotations) >= 1, "fixture_report_2.pdf is known to contain hyperlink annotations"
    for a in annotations:
        assert a.capabilities.editable is False
        assert a.anchor.format == "pdf"


def test_pdf_images_detected_from_scanned_fixture():
    """fixture_report.pdf has no text layer (confirmed separately) but IS
    made of embedded page-raster images — perception must represent that
    as real IMAGE elements, not an empty result, and must not fabricate
    OCR text for it."""
    path = str(FIXTURES / "fixture_report.pdf")
    _, elements = _perceive(path, "pdf")
    images = [e for e in elements if e.type == ElementType.IMAGE]
    assert images, "a scanned PDF must still be represented via detected page images, not silently empty"
    for img in images:
        assert img.source == "text_layer"  # detection metadata, not OCR — text itself is empty, never fabricated
        assert img.text == ""


def test_pdf_write_back_is_honestly_never_editable():
    path = str(FIXTURES / "fixture_report_2.pdf")
    _, elements = _perceive(path, "pdf")
    assert elements
    assert all(e.capabilities.editable is False for e in elements)


# --- Capabilities / element_id stability ------------------------------------


def test_element_id_is_stable_across_reparses_of_an_unchanged_file():
    path = str(FIXTURES / "fixture_generic_handbook.docx")
    _, elements_1 = _perceive(path, "docx")
    _, elements_2 = _perceive(path, "docx")
    ids_1 = [e.element_id for e in elements_1]
    ids_2 = [e.element_id for e in elements_2]
    assert ids_1 == ids_2, "re-perceiving the same unchanged file must yield the same element_id per element"
    assert len(set(ids_1)) == len(ids_1), "element_id must be unique within one document"


def test_editable_capability_matches_actual_writeback_support():
    """DOCX/XLSX text/cell elements are genuinely writeable
    (output/writeback.py handles both); PDF has no writeback handler at
    all (raises ValueError for any non-docx/xlsx format) — capabilities
    must never claim editable=True where output/writeback.py can't
    actually honor a PATCH."""
    _, docx_elements = _perceive(str(FIXTURES / "fixture_generic_handbook.docx"), "docx")
    assert any(e.capabilities.editable for e in docx_elements)

    _, pdf_elements = _perceive(str(FIXTURES / "fixture_report_2.pdf"), "pdf")
    assert not any(e.capabilities.editable for e in pdf_elements)
