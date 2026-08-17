"""Generates fixture_generic_handbook.docx — a diverse, non-financial DOCX
fixture for tests/test_parser_generic.py.

Purpose: every other DOCX fixture in this directory (fixture_bcdt.docx) is a
real financial statement. That leaves "the Geometry Layer is generic, not
financial-statement-shaped" an unverified claim — this fixture and its test
turn it into a checked property, guarding perception/parser.py (and future
perception/anchor_builder.py / perception/element_classifier.py work) against
silently assuming a financial-document layout.

Content is entirely fictional (a community garden handbook) — no real
names, no company names, no financial figures. The structure below (heading
count/levels, paragraph text, table dimensions) is the single source of
truth for both the generated fixture and tests/test_parser_generic.py's
assertions, imported directly from this module rather than duplicated.

Re-run this script to regenerate the fixture after editing the content:
    python tests/fixtures/_generate_generic_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

OUTPUT_PATH = Path(__file__).resolve().parent / "fixture_generic_handbook.docx"

# (heading level, text) — 4 headings across 3 distinct levels.
HEADINGS = [
    (1, "Community Garden Member Handbook"),
    (2, "Getting Started"),
    (2, "Plot Assignments"),
    (3, "Watering Schedule"),
]

# One content paragraph immediately follows each heading above.
PARAGRAPHS = [
    "Welcome to the Riverside Community Garden. This handbook explains how "
    "shared plots are assigned, cared for, and rotated between seasons.",
    "New members should attend an orientation session before claiming a "
    "plot. Orientation covers composting basics, tool storage, and the "
    "shared watering rota.",
    "Plots are assigned on a first-come, first-served basis each spring. "
    "Members who do not renew by the posted deadline forfeit their plot "
    "to the waiting list.",
    "Watering is scheduled in two-hour blocks between 6 a.m. and 8 p.m. "
    "Members should sign up for a slot at the shed noticeboard.",
]

# 3x3 table: row 0 is a header row, rows 1-2 are data.
TABLE_ROWS = [
    ["Plot Size", "Dimensions", "Suggested Crop"],
    ["Small", "4 ft x 4 ft", "Herbs"],
    ["Large", "8 ft x 8 ft", "Squash"],
]


def build_document() -> Document:
    doc = Document()
    for (level, heading_text), paragraph_text in zip(HEADINGS, PARAGRAPHS):
        doc.add_heading(heading_text, level=level)
        doc.add_paragraph(paragraph_text)

    table = doc.add_table(rows=len(TABLE_ROWS), cols=len(TABLE_ROWS[0]))
    for r, row_values in enumerate(TABLE_ROWS):
        for c, value in enumerate(row_values):
            table.rows[r].cells[c].text = value

    return doc


def main() -> None:
    build_document().save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
