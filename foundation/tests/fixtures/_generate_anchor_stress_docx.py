"""Generates fixture_anchor_stress.docx — a synthetic, non-financial DOCX
fixture that reproduces the duplicate-signature ambiguity
`AnchorDOCX.duplicate_ordinal` (perception/anchor_builder.py) exists to
resolve, for tests/test_anchor_p304_synthetic.py.

Purpose: the original P3-04 duplicate_ordinal stress test
(tests/test_anchor_builder.py::test_p304_docx_duplicate_ordinal_survives_uneven_drift_on_real_document)
only runs against a real client DOCX gated by `@requires_real_docx` — on a
clean checkout or in CI without that file, it's silently skipped, leaving
the project's most critical anchor-stability guarantee with zero coverage.
This fixture reproduces the same structural ambiguity synthetically and
non-financially, so the equivalent test can run unconditionally.

Content is entirely fictional (a neighborhood tool-library handbook) — no
client names, no financial figures. The structure below (caption text/style,
repetition count, section content) is the single source of truth for both
the generated fixture and test_anchor_p304_synthetic.py's assertions,
imported directly from this module rather than duplicated.

The key structural feature: CAPTION_TEXT is repeated CAPTION_COUNT times,
every time with the exact same style (CAPTION_STYLE) — the same
(style_id, text_fingerprint) signature `assign_docx_anchor` groups by —
with a heading + a distinct body paragraph between each repetition. That
"other content" between occurrences is what makes uneven drift (inserting
paragraphs between two particular occurrences, not at the very start)
possible to simulate: it shifts only the occurrences after the insertion
point, which is exactly the case a naive "nearest paragraph_index"
tie-break gets wrong and duplicate_ordinal gets right.

Re-run this script to regenerate the fixture after editing the content:
    python tests/fixtures/_generate_anchor_stress_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

OUTPUT_PATH = Path(__file__).resolve().parent / "fixture_anchor_stress.docx"

CAPTION_TEXT = "See the note at the end of this section."
CAPTION_STYLE = "Normal"
# >= 6 required by test_anchor_p304_synthetic.py; a few extra gives a
# comfortable margin around the "middle occurrence" (ordinal 4) the P3-04
# test specifically exercises.
CAPTION_COUNT = 8

TITLE = "Neighborhood Tool Library Handbook"
INTRO = (
    "This handbook explains how to borrow, return, and care for tools from "
    "the neighborhood tool library. It is organized into short sections, "
    "each ending with a pointer to a shared note below."
)

# One (heading, body paragraph) pair per repetition of CAPTION_TEXT —
# len(SECTION_TITLES) == len(SECTION_BODIES) == CAPTION_COUNT.
SECTION_TITLES = [
    "Borrowing a Tool",
    "Returning a Tool",
    "Membership Tiers",
    "Reserving Ahead",
    "Damaged Equipment",
    "Late Fees",
    "Volunteer Shifts",
    "Workshop Signups",
]

SECTION_BODIES = [
    "Members may borrow up to three tools at a time for a period of two weeks.",
    "Tools should be returned clean and in the labeled case they came in.",
    "Standard membership includes access to hand tools; premium membership adds power tools.",
    "Popular tools can be reserved up to one week in advance at the front desk.",
    "Report any damage immediately so the tool can be repaired before the next loan.",
    "A small late fee applies for each day a tool is returned past its due date.",
    "Volunteers keep the library running -- sign up for a shift at the front desk.",
    "Seasonal workshops cover basic tool maintenance and safe handling.",
]

assert len(SECTION_TITLES) == len(SECTION_BODIES) == CAPTION_COUNT

CLOSING_NOTE = (
    "The note referenced throughout this handbook: tool library hours may "
    "vary by season -- check the posted schedule before visiting."
)


def build_document() -> Document:
    doc = Document()
    doc.add_heading(TITLE, level=1)
    doc.add_paragraph(INTRO)

    for section_title, section_body in zip(SECTION_TITLES, SECTION_BODIES):
        doc.add_heading(section_title, level=2)
        doc.add_paragraph(section_body)
        doc.add_paragraph(CAPTION_TEXT, style=CAPTION_STYLE)

    doc.add_heading("Note", level=2)
    doc.add_paragraph(CLOSING_NOTE)

    return doc


def main() -> None:
    build_document().save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
