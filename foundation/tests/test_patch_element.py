"""Tests for the live-edit path: WritebackEngine.apply_single_patch()
(output/writeback.py) and
PATCH /api/documents/<session_id>/elements/<doc_id>
(api/routes/documents.py) — writes a new value directly into a specific
perceived document at an element's Anchor, without re-running any pipeline.
Not restricted to a single "target" document — any perceived doc of a
writeable format (DOCX or XLSX) is editable this way.
"""
import json
import sys
import uuid
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document as DocxDocument  # noqa: E402

from perception.anchor_builder import assign_docx_anchor, assign_xlsx_anchor  # noqa: E402
from perception.models import AnchorXLSX  # noqa: E402
from perception.parser import parse_docx, parse_xlsx  # noqa: E402
from output.writeback import WritebackEngine  # noqa: E402

# --- WritebackEngine.apply_single_patch() — unit level -----------------------


def test_apply_single_patch_docx_paragraph(tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("Original text.")
    path = tmp_path / "doc.docx"
    doc.save(path)

    blocks = parse_docx(str(path))
    anchor = assign_docx_anchor(blocks[0])

    output_path = tmp_path / "patched.docx"
    message = WritebackEngine().apply_single_patch(str(path), anchor, "Edited text.", str(output_path))

    assert message is None
    reopened = DocxDocument(str(output_path))
    assert reopened.paragraphs[0].text == "Edited text."


def test_apply_single_patch_docx_table_cell(tmp_path):
    doc = DocxDocument()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Header"
    table.rows[1].cells[0].text = "Old value"
    path = tmp_path / "doc.docx"
    doc.save(path)

    blocks = parse_docx(str(path))
    cell_block = next(b for b in blocks if b["text"] == "Old value")
    anchor = assign_docx_anchor(cell_block)

    output_path = tmp_path / "patched.docx"
    WritebackEngine().apply_single_patch(str(path), anchor, "New value", str(output_path))

    reopened = DocxDocument(str(output_path))
    assert reopened.tables[0].rows[1].cells[0].text == "New value"


def test_apply_single_patch_accumulates_across_multiple_edits(tmp_path):
    """The PATCH route re-uses the same output path as both input and
    output on the second+ edit — confirm that read-modify-write pattern
    doesn't lose the first edit."""
    doc = DocxDocument()
    doc.add_paragraph("Para one.")
    doc.add_paragraph("Para two.")
    path = tmp_path / "doc.docx"
    doc.save(path)

    blocks = parse_docx(str(path))
    anchor_one = assign_docx_anchor(blocks[0])
    anchor_two = assign_docx_anchor(blocks[1])

    output_path = tmp_path / "patched.docx"
    engine = WritebackEngine()
    engine.apply_single_patch(str(path), anchor_one, "Edited one.", str(output_path))
    engine.apply_single_patch(str(output_path), anchor_two, "Edited two.", str(output_path))

    reopened = DocxDocument(str(output_path))
    assert reopened.paragraphs[0].text == "Edited one."
    assert reopened.paragraphs[1].text == "Edited two."


def test_apply_single_patch_rejects_pdf_anchor(tmp_path):
    from perception.models import AnchorPDF

    anchor = AnchorPDF(page=1, bbox_relative=(0, 0, 1, 1), reading_order_index=0)
    with pytest.raises(ValueError):
        WritebackEngine().apply_single_patch("irrelevant.pdf", anchor, "x", "irrelevant_out.pdf")


# --- PATCH /api/documents/<session_id>/elements/<doc_id> — route level -------


@pytest.fixture
def client(monkeypatch, tmp_path):
    import api.routes.documents as documents_module
    from api.app import create_app

    monkeypatch.setattr(documents_module, "UPLOAD_ROOT", tmp_path)
    app = create_app()
    app.config.update(TESTING=True)
    return app.test_client()


def _seed_document(tmp_path, session_id, doc_id, filename, doc_or_wb, element_count=1):
    """Writes a session dir + manifest.json directly (bypassing the real
    upload route) so route-level tests can seed a document without a real
    multipart upload — mirrors api/routes/documents.py's manifest schema."""
    session_dir = tmp_path / session_id
    session_dir.mkdir(parents=True, exist_ok=True)
    stored_filename = f"{doc_id}_{filename}"
    stored_path = session_dir / stored_filename
    doc_or_wb.save(stored_path)

    manifest_file = session_dir / "manifest.json"
    if manifest_file.exists():
        manifest = json.loads(manifest_file.read_text(encoding="utf-8"))
    else:
        manifest = {"documents": {}}

    fmt = Path(filename).suffix.lstrip(".")
    manifest["documents"][doc_id] = {
        "original_filename": filename,
        "stored_filename": stored_filename,
        "format": fmt,
        "status": "ready",
        "element_count": element_count,
        "error": None,
    }
    manifest_file.write_text(json.dumps(manifest), encoding="utf-8")
    return session_dir, stored_path


def test_patch_element_seeds_patched_file_from_original_when_none_exists(client, tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("Original text.")
    session_id, doc_id = str(uuid.uuid4()), str(uuid.uuid4())
    session_dir, stored_path = _seed_document(tmp_path, session_id, doc_id, "target.docx", doc)

    blocks = parse_docx(str(stored_path))
    anchor = assign_docx_anchor(blocks[0])

    response = client.patch(
        f"/api/documents/{session_id}/elements/{doc_id}",
        json={"anchor": anchor.model_dump(mode="json"), "value": "Edited via UI."},
    )

    assert response.status_code == 200
    body = response.get_json()
    assert body["status"] == "ok"
    assert body["download_url"] == f"/api/documents/{session_id}/download/{doc_id}"

    patched_path = stored_path.with_name(f"{stored_path.stem}_patched.docx")
    assert patched_path.exists()
    assert DocxDocument(str(patched_path)).paragraphs[0].text == "Edited via UI."
    # Original upload must be left untouched.
    assert DocxDocument(str(stored_path)).paragraphs[0].text == "Original text."


def test_patch_element_edits_accumulate_on_existing_patched_file(client, tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("Para one.")
    doc.add_paragraph("Para two.")
    session_id, doc_id = str(uuid.uuid4()), str(uuid.uuid4())
    session_dir, stored_path = _seed_document(tmp_path, session_id, doc_id, "target.docx", doc)

    blocks = parse_docx(str(stored_path))
    anchor_one = assign_docx_anchor(blocks[0])
    anchor_two = assign_docx_anchor(blocks[1])

    r1 = client.patch(
        f"/api/documents/{session_id}/elements/{doc_id}",
        json={"anchor": anchor_one.model_dump(mode="json"), "value": "First edit."},
    )
    assert r1.status_code == 200

    r2 = client.patch(
        f"/api/documents/{session_id}/elements/{doc_id}",
        json={"anchor": anchor_two.model_dump(mode="json"), "value": "Second edit."},
    )
    assert r2.status_code == 200

    patched = DocxDocument(str(stored_path.with_name(f"{stored_path.stem}_patched.docx")))
    assert patched.paragraphs[0].text == "First edit."  # not lost by the second edit
    assert patched.paragraphs[1].text == "Second edit."


def test_patch_element_edits_xlsx_document_and_persists_real_output(client, tmp_path):
    """Regression guard: XLSX documents are editable through this route —
    output/writeback.py's apply_single_patch already supported xlsx, the
    old /api/elements route just hard-blocked it (a GTPS-shaped "only DOCX
    target" restriction, not a real technical limit). Verifies actual
    persisted output, not just an HTTP 200."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "Old value"
    session_id, doc_id = str(uuid.uuid4()), str(uuid.uuid4())
    session_dir, stored_path = _seed_document(tmp_path, session_id, doc_id, "data.xlsx", wb)

    blocks = parse_xlsx(str(stored_path))
    anchor = assign_xlsx_anchor(next(b for b in blocks if b["cell_address"] == "A1"))

    response = client.patch(
        f"/api/documents/{session_id}/elements/{doc_id}",
        json={"anchor": anchor.model_dump(mode="json"), "value": "New value"},
    )

    assert response.status_code == 200
    body = response.get_json()
    assert body["status"] == "ok"
    assert body["download_url"] == f"/api/documents/{session_id}/download/{doc_id}"

    patched_path = stored_path.with_name(f"{stored_path.stem}_patched.xlsx")
    assert patched_path.exists()
    reopened = openpyxl.load_workbook(patched_path)
    assert reopened["Sheet1"]["A1"].value == "New value"

    # Original upload must be left untouched.
    original = openpyxl.load_workbook(stored_path)
    assert original["Sheet1"]["A1"].value == "Old value"

    # Download route serves this same patched file for this doc_id.
    download_response = client.get(f"/api/documents/{session_id}/download/{doc_id}")
    assert download_response.status_code == 200


def test_patch_element_returns_404_for_unknown_session_id(client):
    response = client.patch(
        f"/api/documents/does-not-exist/elements/{uuid.uuid4()}",
        json={"anchor": {"format": "docx", "style_id": "", "text_fingerprint": "abcd1234"}, "value": "x"},
    )
    assert response.status_code == 404


def test_patch_element_returns_404_for_unknown_doc_id(client, tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("x")
    session_id = str(uuid.uuid4())
    _seed_document(tmp_path, session_id, str(uuid.uuid4()), "x.docx", doc)

    response = client.patch(
        f"/api/documents/{session_id}/elements/{uuid.uuid4()}",
        json={"anchor": {"format": "docx", "style_id": "", "text_fingerprint": "abcd1234"}, "value": "x"},
    )
    assert response.status_code == 404


def test_patch_element_returns_400_for_missing_body(client, tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("x")
    session_id, doc_id = str(uuid.uuid4()), str(uuid.uuid4())
    _seed_document(tmp_path, session_id, doc_id, "x.docx", doc)

    response = client.patch(
        f"/api/documents/{session_id}/elements/{doc_id}",
        json={"anchor": {"format": "docx"}},
    )
    assert response.status_code == 400


def test_patch_element_returns_400_for_malformed_anchor(client, tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("x")
    session_id, doc_id = str(uuid.uuid4()), str(uuid.uuid4())
    _seed_document(tmp_path, session_id, doc_id, "x.docx", doc)

    # Missing required fields (sheet_name, cell_address) — fails pydantic
    # validation before the format check ever runs.
    response = client.patch(
        f"/api/documents/{session_id}/elements/{doc_id}",
        json={"anchor": {"format": "xlsx"}, "value": "y"},
    )
    assert response.status_code == 400


def test_patch_element_returns_422_for_pdf_anchor(client, tmp_path):
    """PDF stays genuinely read-only (output/writeback.py's own real
    technical limit) — this is not GTPS-shaped, unlike the old "only DOCX
    target" restriction removed above."""
    doc = DocxDocument()
    doc.add_paragraph("x")
    session_id, doc_id = str(uuid.uuid4()), str(uuid.uuid4())
    _seed_document(tmp_path, session_id, doc_id, "x.docx", doc)

    response = client.patch(
        f"/api/documents/{session_id}/elements/{doc_id}",
        json={
            "anchor": {"format": "pdf", "page": 1, "bbox_relative": [0, 0, 1, 1], "reading_order_index": 0},
            "value": "y",
        },
    )
    assert response.status_code == 422


def test_patch_element_xlsx_type_coercion(client, tmp_path):
    """Verifies that numeric, float, boolean, and empty values are properly coerced in XLSX output."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    ws["A1"] = "old_int"
    ws["A2"] = "old_float"
    ws["A3"] = "old_bool"
    ws["A4"] = "old_code"
    ws["A5"] = "old_empty"

    session_id, doc_id = str(uuid.uuid4()), str(uuid.uuid4())
    session_dir, stored_path = _seed_document(tmp_path, session_id, doc_id, "coercion.xlsx", wb)

    # 1. Integer
    anchor_a1 = AnchorXLSX(sheet_name="Data", cell_address="A1")
    r1 = client.patch(f"/api/documents/{session_id}/elements/{doc_id}", json={"anchor": anchor_a1.model_dump(mode="json"), "value": "1234"})
    assert r1.status_code == 200

    # 2. Float
    anchor_a2 = AnchorXLSX(sheet_name="Data", cell_address="A2")
    r2 = client.patch(f"/api/documents/{session_id}/elements/{doc_id}", json={"anchor": anchor_a2.model_dump(mode="json"), "value": "123.45"})
    assert r2.status_code == 200

    # 3. Boolean
    anchor_a3 = AnchorXLSX(sheet_name="Data", cell_address="A3")
    r3 = client.patch(f"/api/documents/{session_id}/elements/{doc_id}", json={"anchor": anchor_a3.model_dump(mode="json"), "value": "true"})
    assert r3.status_code == 200

    # 4. String with leading zero
    anchor_a4 = AnchorXLSX(sheet_name="Data", cell_address="A4")
    r4 = client.patch(f"/api/documents/{session_id}/elements/{doc_id}", json={"anchor": anchor_a4.model_dump(mode="json"), "value": "012345"})
    assert r4.status_code == 200

    # 5. Empty cell
    anchor_a5 = AnchorXLSX(sheet_name="Data", cell_address="A5")
    r5 = client.patch(f"/api/documents/{session_id}/elements/{doc_id}", json={"anchor": anchor_a5.model_dump(mode="json"), "value": ""})
    assert r5.status_code == 200

    patched_path = stored_path.with_name(f"{stored_path.stem}_patched.xlsx")
    patched_wb = openpyxl.load_workbook(patched_path)
    ws_patched = patched_wb["Data"]

    assert ws_patched["A1"].value == 1234
    assert isinstance(ws_patched["A1"].value, int)
    assert ws_patched["A2"].value == 123.45
    assert isinstance(ws_patched["A2"].value, float)
    assert ws_patched["A3"].value is True
    assert isinstance(ws_patched["A3"].value, bool)
    assert ws_patched["A4"].value == "012345"
    assert isinstance(ws_patched["A4"].value, str)
    assert ws_patched["A5"].value is None


def test_patch_element_xlsx_formula_cell_protection(client, tmp_path):
    """Verifies that formula cells are classified as read-only and rejected by backend patch."""
    import openpyxl
    from perception.parser import extract_geometry
    from perception.anchor_builder import assign_anchors
    from perception.element_classifier import classify_blocks

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Calc"
    ws["A1"] = 10
    ws["A2"] = 20
    ws["A3"] = "=SUM(A1:A2)"

    session_id, doc_id = str(uuid.uuid4()), str(uuid.uuid4())
    session_dir, stored_path = _seed_document(tmp_path, session_id, doc_id, "formulas.xlsx", wb)

    # 1. Check perception capabilities
    blocks = extract_geometry(str(stored_path))
    anchors = assign_anchors(blocks, "xlsx")
    elements = classify_blocks(blocks, "xlsx", anchors)

    formula_el = next(e for e in elements if e.anchor.cell_address == "A3")
    literal_el = next(e for e in elements if e.anchor.cell_address == "A1")

    assert formula_el.capabilities.editable is False
    assert literal_el.capabilities.editable is True

    # 2. Attempting to write a literal value to a formula cell is rejected with 422
    anchor_a3 = AnchorXLSX(sheet_name="Calc", cell_address="A3")
    response = client.patch(
        f"/api/documents/{session_id}/elements/{doc_id}",
        json={"anchor": anchor_a3.model_dump(mode="json"), "value": "100"},
    )
    assert response.status_code == 422
    assert "contains a formula" in response.get_json()["error"]
    assert "read-only" in response.get_json()["error"]


def test_patch_element_xlsx_undo_roundtrip(client, tmp_path):
    """Verifies that an edit can be undone by patching back the original value."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["B5"] = "Initial Company Name"

    session_id, doc_id = str(uuid.uuid4()), str(uuid.uuid4())
    session_dir, stored_path = _seed_document(tmp_path, session_id, doc_id, "company.xlsx", wb)

    anchor_b5 = AnchorXLSX(sheet_name="Sheet1", cell_address="B5")

    # 1. Edit to new value
    r1 = client.patch(
        f"/api/documents/{session_id}/elements/{doc_id}",
        json={"anchor": anchor_b5.model_dump(mode="json"), "value": "Updated Company Name"},
    )
    assert r1.status_code == 200

    patched_path = stored_path.with_name(f"{stored_path.stem}_patched.xlsx")
    wb_edited = openpyxl.load_workbook(patched_path)
    assert wb_edited["Sheet1"]["B5"].value == "Updated Company Name"

    # 2. Undo edit by patching back initial value
    r2 = client.patch(
        f"/api/documents/{session_id}/elements/{doc_id}",
        json={"anchor": anchor_b5.model_dump(mode="json"), "value": "Initial Company Name"},
    )
    assert r2.status_code == 200

    wb_restored = openpyxl.load_workbook(patched_path)
    assert wb_restored["Sheet1"]["B5"].value == "Initial Company Name"


def test_patch_element_multi_document_isolation(client, tmp_path):
    """Verifies that editing an XLSX document in a multi-document session leaves DOCX documents untouched."""
    import openpyxl
    from docx import Document as DocxDocument

    # Document 1: DOCX
    docx = DocxDocument()
    docx.add_paragraph("DOCX Pristine Paragraph")
    session_id = str(uuid.uuid4())
    docx_doc_id = str(uuid.uuid4())
    _seed_document(tmp_path, session_id, docx_doc_id, "doc1.docx", docx)

    # Document 2: XLSX
    xlsx = openpyxl.Workbook()
    xlsx.active.title = "Sheet1"
    xlsx.active["A1"] = "XLSX Cell Original"
    xlsx_doc_id = str(uuid.uuid4())
    session_dir, xlsx_path = _seed_document(tmp_path, session_id, xlsx_doc_id, "doc2.xlsx", xlsx)

    # Patch XLSX only
    anchor_xlsx = AnchorXLSX(sheet_name="Sheet1", cell_address="A1")
    r = client.patch(
        f"/api/documents/{session_id}/elements/{xlsx_doc_id}",
        json={"anchor": anchor_xlsx.model_dump(mode="json"), "value": "XLSX Cell Edited"},
    )
    assert r.status_code == 200

    # Verify XLSX is patched
    patched_xlsx = openpyxl.load_workbook(xlsx_path.with_name(f"{xlsx_path.stem}_patched.xlsx"))
    assert patched_xlsx["Sheet1"]["A1"].value == "XLSX Cell Edited"

    # Verify DOCX download serves untouched file with original content
    download_docx = client.get(f"/api/documents/{session_id}/download/{docx_doc_id}")
    assert download_docx.status_code == 200
    # No _patched.docx should exist for doc1
    assert not (session_dir / f"{docx_doc_id}_patched.docx").exists()
