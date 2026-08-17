"""Unit tests for perception/anchor_builder.py — the Locate capability
(assign + resolve stable Anchors), including the P3-04 acceptance test
(Foundation_Build_Plan_v3.md §9.2 / Foundation_Master_Context.md §5):
insert a paragraph, re-parse, resolve the old anchor, must still find the
same element.
"""
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document as DocxDocument  # noqa: E402
import openpyxl  # noqa: E402

from perception.anchor_builder import (  # noqa: E402
    assign_anchors,
    assign_docx_anchor,
    assign_pdf_anchors,
    assign_xlsx_anchor,
    resolve_docx_anchor,
    resolve_paragraph_anchor,
    resolve_pdf_anchor,
    resolve_xlsx_anchor,
)
from perception.parser import parse_docx, parse_pdf, parse_xlsx  # noqa: E402

FIXTURES = Path(__file__).resolve().parent / "fixtures"

# Real HMV Local File — same fixture test_mapping_service.py uses. Genuinely
# has repeated captions (e.g. "Source: TP Cat database" appears 9 times,
# same BodyText style) — exactly the ambiguity duplicate_ordinal exists for.
_REAL_DOCX = (
    Path(__file__).resolve().parents[2]
    / "anonymize client"
    / "Demo files"
    / "Demo files"
    / "Compare LF"
    / "HMV-26-Final-Local File for FY2024-EN-R2901KPMG_drifted.docx"
)
requires_real_docx = pytest.mark.skipif(
    not _REAL_DOCX.exists(), reason="Real HMV demo fixture not present on this machine"
)


def _pdf_block(text, page, bbox, page_width=612.0, page_height=792.0):
    return {
        "text": text,
        "paragraph_index": None,
        "style_id": None,
        "table_index": None,
        "table_hash": None,
        "row_index": None,
        "col_index": None,
        "page": page,
        "bbox": bbox,
        "page_width": page_width,
        "page_height": page_height,
        "sheet_name": None,
        "cell_address": None,
        "named_range": None,
    }


def _insert_paragraph_before(doc, index: int, text: str):
    """Insert a new paragraph immediately before doc.paragraphs[index]."""
    anchor_para = doc.paragraphs[index]
    new_para = doc.add_paragraph(text)  # appends at the end
    new_para._p.getparent().remove(new_para._p)
    anchor_para._p.addprevious(new_para._p)


# --- Assign: sanity checks ---------------------------------------------------


def test_assign_docx_anchor_paragraph(tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("A Heading", style="Heading 1")
    path = tmp_path / "doc.docx"
    doc.save(path)

    blocks = parse_docx(str(path))
    anchor = assign_docx_anchor(blocks[0])

    assert anchor.format == "docx"
    assert anchor.paragraph_index == 0
    assert anchor.table_index is None
    assert len(anchor.text_fingerprint) == 8


def test_assign_docx_anchor_table_cell(tmp_path):
    doc = DocxDocument()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Header"
    table.rows[1].cells[0].text = "Value"
    path = tmp_path / "doc.docx"
    doc.save(path)

    blocks = parse_docx(str(path))
    cell_block = next(b for b in blocks if b["table_index"] is not None)
    anchor = assign_docx_anchor(cell_block)

    assert anchor.paragraph_index is None
    assert anchor.table_index == 0
    assert anchor.table_hash is not None
    assert anchor.row_index == cell_block["row_index"]
    assert anchor.col_index == cell_block["col_index"]


def test_assign_xlsx_anchor_with_and_without_named_range(tmp_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "Named"
    ws["B2"] = "Plain"
    wb.defined_names["MyRange"] = openpyxl.workbook.defined_name.DefinedName(
        "MyRange", attr_text="Sheet1!$A$1"
    )
    path = tmp_path / "wb.xlsx"
    wb.save(path)
    wb.close()

    blocks = parse_xlsx(str(path))
    named_block = next(b for b in blocks if b["cell_address"] == "A1")
    plain_block = next(b for b in blocks if b["cell_address"] == "B2")

    named_anchor = assign_xlsx_anchor(named_block)
    plain_anchor = assign_xlsx_anchor(plain_block)

    assert named_anchor.named_range == "MyRange"
    assert plain_anchor.named_range is None
    assert plain_anchor.cell_address == "B2"


def test_assign_pdf_anchors_reading_order_resets_per_page():
    blocks = [
        _pdf_block("line1", page=1, bbox=(10, 10, 100, 20)),
        _pdf_block("line2", page=1, bbox=(10, 30, 100, 40)),
        _pdf_block("line3", page=2, bbox=(10, 10, 100, 20)),
    ]
    anchors = assign_pdf_anchors(blocks)

    assert [a.reading_order_index for a in anchors] == [0, 1, 0]
    assert anchors[0].page == 1
    assert anchors[2].page == 2
    for a in anchors:
        for v in a.bbox_relative:
            assert 0.0 <= v <= 1.0


def test_assign_anchors_dispatches_by_format():
    blocks = [_pdf_block("x", page=1, bbox=(0, 0, 10, 10))]
    anchors = assign_anchors(blocks, "pdf")
    assert anchors[0].format == "pdf"

    with pytest.raises(ValueError):
        assign_anchors(blocks, "unknown")


# --- P3-04: DOCX paragraph anchor stability ----------------------------------


def test_p304_docx_paragraph_anchor_survives_insertion_at_start(tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("Intro", style="Heading 1")
    doc.add_paragraph("Body text one.")
    doc.add_paragraph("Body text two.")
    path_v1 = tmp_path / "v1.docx"
    doc.save(path_v1)

    blocks_v1 = parse_docx(str(path_v1))
    target_block = next(b for b in blocks_v1 if b["text"] == "Body text one.")
    anchor = assign_docx_anchor(target_block)
    assert anchor.paragraph_index == 1  # before insertion

    # Simulate structural drift: insert a new paragraph at the very start —
    # this is exactly the P3-04 acceptance scenario.
    doc_v2 = DocxDocument(str(path_v1))
    _insert_paragraph_before(doc_v2, 0, "A brand new intro paragraph.")
    path_v2 = tmp_path / "v2.docx"
    doc_v2.save(path_v2)

    doc_for_resolve = DocxDocument(str(path_v2))
    resolved, message = resolve_docx_anchor(doc_for_resolve, anchor)

    assert resolved.text == "Body text one."
    assert message is None  # Strategy 1 — clean content match, no warning


@requires_real_docx
def test_p304_docx_duplicate_ordinal_survives_uneven_drift_on_real_document(tmp_path):
    """Real-world stress test for Strategy 1 tie-breaking. The real HMV
    Local File repeats the caption "Source: TP Cat database" 9 times
    (same BodyText style) — genuine duplicate-signature ambiguity, not a
    synthetic worst case.

    Critically, this test inserts paragraphs *between* two occurrences
    (not at the very start), which shifts only the occurrences after the
    insertion point. A tie-break based on "nearest to the originally
    recorded paragraph_index" gets this wrong — the now-much-closer
    *neighboring* occurrence wins instead of the actual target. This is
    exactly the failure duplicate_ordinal (rank among same-signature
    paragraphs, immune to uneven shifts) exists to fix.
    """
    blocks = parse_docx(str(_REAL_DOCX))
    occurrences = [b for b in blocks if b["text"].startswith("Source: TP Cat database")]
    assert len(occurrences) >= 5, "fixture no longer has the expected repeated caption"

    target_block = occurrences[4]  # 5th occurrence (0-indexed ordinal 4)
    target_index = target_block["paragraph_index"]

    anchors = assign_anchors(blocks, "docx")
    anchor = anchors[blocks.index(target_block)]
    assert anchor.duplicate_ordinal == 4

    # Insert 50 unrelated paragraphs immediately before the target — this
    # shifts the target (and everything after it) by +50, while the
    # *previous* occurrence (3 paragraphs earlier, per the fixture's
    # regular spacing) stays completely unmoved and becomes the closer
    # match by raw paragraph_index distance.
    doc = DocxDocument(str(_REAL_DOCX))
    for i in range(50):
        _insert_paragraph_before(doc, target_index, f"Injected filler paragraph {i}.")
    drifted_path = tmp_path / "drifted.docx"
    doc.save(drifted_path)

    doc_for_resolve = DocxDocument(str(drifted_path))
    expected_target = doc_for_resolve.paragraphs[target_index + 50]
    assert expected_target.text == "Source: TP Cat database"

    resolved, message = resolve_docx_anchor(doc_for_resolve, anchor)

    assert resolved.text == "Source: TP Cat database"
    assert resolved._p is expected_target._p, (
        "resolved the wrong occurrence of a duplicated caption — likely "
        "picked the nearer, unrelated neighbor instead of the actual target"
    )
    assert message is None  # Strategy 1, ordinal-disambiguated — no warning


def test_docx_paragraph_resolve_falls_back_to_strategy_2_when_text_changed(tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("Intro", style="Heading 1")
    doc.add_paragraph("Original text.")
    path_v1 = tmp_path / "v1.docx"
    doc.save(path_v1)

    blocks_v1 = parse_docx(str(path_v1))
    target_block = next(b for b in blocks_v1 if b["text"] == "Original text.")
    anchor = assign_docx_anchor(target_block)

    # Edit the target paragraph's text in place (style + position unchanged,
    # so the text_fingerprint no longer matches — Strategy 1 must miss).
    doc_v2 = DocxDocument(str(path_v1))
    doc_v2.paragraphs[1].runs[0].text = "Edited text, different content."
    path_v2 = tmp_path / "v2.docx"
    doc_v2.save(path_v2)

    resolved, message = resolve_docx_anchor(DocxDocument(str(path_v2)), anchor)

    assert resolved.text == "Edited text, different content."
    assert message is not None and "Strategy 2" in message


def test_docx_paragraph_resolve_raises_when_anchor_points_past_end(tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("Only paragraph.")
    path = tmp_path / "doc.docx"
    doc.save(path)

    blocks = parse_docx(str(path))
    anchor = assign_docx_anchor(blocks[0])
    anchor.paragraph_index = 99
    anchor.text_fingerprint = "ffffffff"  # force a guaranteed non-match

    # resolve_paragraph_anchor itself returns (None, None) on total failure —
    # resolve_docx_anchor is the layer that raises (never a silent miss).
    result, message = resolve_paragraph_anchor(doc, anchor)
    assert result is None and message is None

    with pytest.raises(ValueError):
        resolve_docx_anchor(DocxDocument(str(path)), anchor)


# --- XLSX resolve -------------------------------------------------------------


def test_xlsx_resolve_by_cell_address_and_named_range(tmp_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "Named value"
    ws["B2"] = "Plain value"
    wb.defined_names["Revenue"] = openpyxl.workbook.defined_name.DefinedName(
        "Revenue", attr_text="Sheet1!$A$1"
    )
    path = tmp_path / "wb.xlsx"
    wb.save(path)
    wb.close()

    blocks = parse_xlsx(str(path))
    named_block = next(b for b in blocks if b["cell_address"] == "A1")
    plain_block = next(b for b in blocks if b["cell_address"] == "B2")
    named_anchor = assign_xlsx_anchor(named_block)
    plain_anchor = assign_xlsx_anchor(plain_block)

    wb2 = openpyxl.load_workbook(path)
    named_cell, named_message = resolve_xlsx_anchor(wb2, named_anchor)
    plain_cell, plain_message = resolve_xlsx_anchor(wb2, plain_anchor)
    assert named_cell.value == "Named value" and named_message is None
    assert plain_cell.value == "Plain value" and plain_message is None


def test_xlsx_resolve_self_heals_when_row_inserted_above(tmp_path):
    """The scenario STATUS.md's original "cell_address is stable enough"
    decision didn't cover: a row inserted above the target silently makes
    a plain cell_address point at the wrong row. row_label_fingerprint
    (the row's line-item label, e.g. "Net profit" in a financial
    statement) should catch this and self-heal — same idea as
    resolve_table_anchor's table-hash heal for DOCX."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "FA"
    ws["A1"] = "Line Item"
    ws["B1"] = "FY2024"
    ws["A2"] = "Revenue"
    ws["B2"] = 1000
    ws["A3"] = "COGS"
    ws["B3"] = 400
    ws["A4"] = "Net profit"
    ws["B4"] = 600
    path = tmp_path / "fs.xlsx"
    wb.save(path)

    blocks = parse_xlsx(str(path))
    target_block = next(b for b in blocks if b["cell_address"] == "B4")
    assert target_block["row_label"] == "Net profit"
    anchor = assign_xlsx_anchor(target_block)
    assert anchor.row_label_fingerprint is not None

    # Insert a new line item above "Net profit" — B4 now silently points
    # at the wrong row (the newly inserted one) if resolved naively.
    wb2 = openpyxl.load_workbook(path)
    ws2 = wb2["FA"]
    ws2.insert_rows(3)
    ws2["A3"] = "New line item"
    ws2["B3"] = 999
    wb2.save(path)

    wb3 = openpyxl.load_workbook(path)
    cell, message = resolve_xlsx_anchor(wb3, anchor)

    assert cell.value == 600  # the actual "Net profit" value, not 999
    assert cell.coordinate == "B5"
    assert message is not None and "Self-healed" in message


def test_xlsx_resolve_raises_when_row_label_gone(tmp_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "FA"
    ws["A1"] = "Revenue"
    ws["B1"] = 1000
    path = tmp_path / "fs.xlsx"
    wb.save(path)

    blocks = parse_xlsx(str(path))
    anchor = assign_xlsx_anchor(next(b for b in blocks if b["cell_address"] == "B1"))

    wb2 = openpyxl.load_workbook(path)
    ws2 = wb2["FA"]
    ws2["A1"] = "Completely different label"
    ws2["B1"] = 2000
    wb2.save(path)

    wb3 = openpyxl.load_workbook(path)
    with pytest.raises(ValueError):
        resolve_xlsx_anchor(wb3, anchor)


# --- PDF resolve ---------------------------------------------------------------


def test_pdf_resolve_strategy1_exact_reading_order_match():
    blocks = [
        _pdf_block("header", page=1, bbox=(10, 10, 100, 20)),
        _pdf_block("target line", page=1, bbox=(10, 30, 100, 40)),
    ]
    anchor = assign_pdf_anchors(blocks)[1]

    resolved, message = resolve_pdf_anchor(blocks, anchor)

    assert resolved["text"] == "target line"
    assert message is None


def test_pdf_resolve_strategy2_bbox_fallback_when_reading_order_shifts():
    blocks_v1 = [
        _pdf_block("target line", page=1, bbox=(10, 30, 100, 40)),
    ]
    anchor = assign_pdf_anchors(blocks_v1)[0]
    assert anchor.reading_order_index == 0

    # A new line was inserted above it — same bbox for our target, but its
    # reading_order_index within the page has shifted from 0 to 1.
    blocks_v2 = [
        _pdf_block("a new inserted line", page=1, bbox=(10, 5, 100, 15)),
        _pdf_block("target line", page=1, bbox=(10, 30, 100, 40)),
    ]

    resolved, message = resolve_pdf_anchor(blocks_v2, anchor)

    assert resolved["text"] == "target line"
    assert message is not None and "Strategy 2" in message


def test_pdf_resolve_raises_when_page_missing():
    anchor = assign_pdf_anchors([_pdf_block("x", page=1, bbox=(0, 0, 10, 10))])[0]
    with pytest.raises(ValueError):
        resolve_pdf_anchor([_pdf_block("y", page=2, bbox=(0, 0, 10, 10))], anchor)


_REAL_PDF = FIXTURES / "fixture_report_2.pdf"
requires_real_pdf = pytest.mark.skipif(
    not _REAL_PDF.exists(), reason="Real 32-page PDF fixture not present on this machine"
)


@requires_real_pdf
def test_pdf_resolve_survives_realistic_insertion_on_real_dense_document():
    """Real-data stress test, not fully synthetic: uses the actual parsed
    blocks of a genuine 32-page financial statement (968 lines total,
    including real repeated boilerplate across many pages — see
    tests/test_parser.py's diversity test on this same fixture). Simulates
    a realistic edit (one new line inserted on one page) directly on the
    real block list, then confirms resolve still finds the exact original
    line and doesn't get confused by the document's genuine duplicate-text
    noise on other pages.
    """
    blocks = parse_pdf(str(_REAL_PDF))

    page5 = [b for b in blocks if b["page"] == 5]
    target_block = page5[8]  # "Báo cáo kiểm toán số: ..." — a real, distinctive line
    assert blocks.count(target_block) == 1  # not accidentally duplicated content

    anchor = assign_pdf_anchors(blocks)[blocks.index(target_block)]

    # Simulate a realistic edit: one new line inserted on page 5, right
    # before the target — every other page's real content (including its
    # genuine repeated boilerplate) is left completely untouched.
    # Target's real bbox top is ~308pt; place the inserted line far enough
    # away (>0.02 relative page-height tolerance, ~17pt on this page) that
    # Strategy 1 can't accidentally "confirm" the wrong (inserted) block
    # just because it now sits at the old reading_order_index — this is
    # the exact ambiguity Strategy 1's bbox corroboration exists to catch.
    target_pos = blocks.index(target_block)
    inserted_block = _pdf_block(
        "A newly added disclosure line.",
        page=5,
        bbox=(48.54, 150.0, 300.0, 160.0),
        page_width=target_block["page_width"],
        page_height=target_block["page_height"],
    )
    drifted_blocks = blocks[:target_pos] + [inserted_block] + blocks[target_pos:]

    resolved, message = resolve_pdf_anchor(drifted_blocks, anchor)

    assert resolved["text"] == target_block["text"]
    assert resolved is not inserted_block
    assert message is not None and "Strategy 2" in message
