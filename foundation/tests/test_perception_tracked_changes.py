"""Tests for `<w:ins>` (tracked-change insertion) text extraction —
perception/parser.py::_paragraph_text.

python-docx's `Paragraph.text` only sees `<w:r>` elements that are DIRECT
children of `<w:p>`; a tracked insertion wraps its run(s) one level deeper
(`<w:p><w:ins><w:r>...</w:r></w:ins></w:p>`), so that text is invisible to
python-docx entirely — not degraded, genuinely absent. This was confirmed
against the real KPMG fixture before writing the fix (see conversation):
one paragraph, entirely insertion-wrapped, had python-docx report an empty
`.text` and was silently dropped by `parse_docx`'s `if text.strip():` check.

`<w:del>` is deliberately out of scope here: docx-preview's default render
(this project never sets `renderChanges: true`) shows `<w:ins>` content as
plain text and `<w:del>` content as nothing at all — an "accepted changes"
view. Perception text should match what's actually rendered, so deleted
text staying invisible to extraction is correct, not a gap.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from perception.anchor_builder import assign_anchors
from perception.element_classifier import classify_blocks
from perception.parser import parse_docx

_NEXT_INS_ID = [1000]


def _add_inserted_run(paragraph, text: str, author: str = "Reviewer"):
    """Appends a `<w:ins>` — wrapping one `<w:r><w:t>` — as the next child
    of `paragraph._p`. python-docx has no high-level API for tracked
    insertions (they're a review-workflow construct, not part of its
    "build a document" surface), so this manipulates the underlying OOXML
    directly, the same direct-inspection approach already used elsewhere
    in this project for object kinds python-docx doesn't expose."""
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(_NEXT_INS_ID[0]))
    _NEXT_INS_ID[0] += 1
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), "2026-01-01T00:00:00Z")

    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    ins.append(run)
    paragraph._p.append(ins)


def _perceive_docx(path: str):
    blocks = parse_docx(path)
    anchors = assign_anchors(blocks, "docx")
    return blocks, classify_blocks(blocks, "docx", anchors)


def test_normal_paragraph_unaffected(tmp_path):
    """Regression control: a plain paragraph with no tracked changes must
    extract exactly as before — the `<w:ins>` fix must not touch the
    ordinary run-text path at all."""
    doc = DocxDocument()
    doc.add_paragraph("Plain paragraph, no tracked changes.")
    path = tmp_path / "doc.docx"
    doc.save(path)

    blocks = parse_docx(str(path))
    para_blocks = [b for b in blocks if b["kind"] == "paragraph"]
    assert len(para_blocks) == 1
    assert para_blocks[0]["text"] == "Plain paragraph, no tracked changes."
    assert para_blocks[0]["extra"] is None


def test_paragraph_entirely_inserted_text_is_extracted(tmp_path):
    """The exact real-fixture case: a paragraph whose ONLY content is a
    tracked insertion. Before this fix, python-docx's `.text` was empty
    and the paragraph was dropped entirely — not merely truncated."""
    doc = DocxDocument()
    p = doc.add_paragraph()
    _add_inserted_run(p, "Fully inserted paragraph text.")
    path = tmp_path / "doc.docx"
    doc.save(path)

    blocks = parse_docx(str(path))
    para_blocks = [b for b in blocks if b["kind"] == "paragraph"]
    assert len(para_blocks) == 1, "an entirely-inserted paragraph must still become a block, not be silently dropped"
    assert para_blocks[0]["text"] == "Fully inserted paragraph text."
    assert para_blocks[0]["extra"] == {"has_tracked_insertion": True}


def test_inserted_text_mixed_with_normal_text_interleaves_in_order(tmp_path):
    """Normal and inserted runs in the same paragraph must combine in
    DOCUMENT ORDER, matching what docx-preview actually renders — not
    "normal text, then all insertions appended at the end"."""
    doc = DocxDocument()
    p = doc.add_paragraph("Before. ")
    _add_inserted_run(p, "Inserted middle. ")
    p.add_run("After.")
    path = tmp_path / "doc.docx"
    doc.save(path)

    blocks = parse_docx(str(path))
    para_blocks = [b for b in blocks if b["kind"] == "paragraph"]
    assert len(para_blocks) == 1
    assert para_blocks[0]["text"] == "Before. Inserted middle. After."
    assert para_blocks[0]["extra"] == {"has_tracked_insertion": True}


def test_multiple_inserted_runs_all_captured(tmp_path):
    """More than one `<w:ins>` block in the same paragraph — each must
    contribute its text, in order, not just the first/last one."""
    doc = DocxDocument()
    p = doc.add_paragraph("Start. ")
    _add_inserted_run(p, "First insertion. ")
    p.add_run("Middle. ")
    _add_inserted_run(p, "Second insertion.")
    path = tmp_path / "doc.docx"
    doc.save(path)

    blocks = parse_docx(str(path))
    para_blocks = [b for b in blocks if b["kind"] == "paragraph"]
    assert len(para_blocks) == 1
    assert para_blocks[0]["text"] == "Start. First insertion. Middle. Second insertion."


def test_duplicate_inserted_text_gets_distinct_element_ids(tmp_path):
    """Two paragraphs with byte-identical inserted content must still
    resolve to two distinct, stable element_ids — duplicate_ordinal
    disambiguation (anchor_builder.py) must keep working with `<w:ins>`
    content exactly as it already does for plain-text duplicates."""
    doc = DocxDocument()
    p1 = doc.add_paragraph()
    _add_inserted_run(p1, "Repeated inserted boilerplate.")
    doc.add_paragraph("Something in between.")
    p2 = doc.add_paragraph()
    _add_inserted_run(p2, "Repeated inserted boilerplate.")
    path = tmp_path / "doc.docx"
    doc.save(path)

    _, elements = _perceive_docx(str(path))
    matches = [e for e in elements if e.text == "Repeated inserted boilerplate."]
    assert len(matches) == 2
    assert matches[0].element_id != matches[1].element_id
    assert {e.anchor.duplicate_ordinal for e in matches} == {0, 1}

    # Determinism: re-running perception on the same bytes must reproduce
    # the exact same element_ids (element_id is a pure hash of anchor
    # content — no randomness, no dependence on parse order beyond what
    # duplicate_ordinal already accounts for).
    _, elements_again = _perceive_docx(str(path))
    ids_first = [e.element_id for e in elements_again if e.text == "Repeated inserted boilerplate."]
    ids_second = [e.element_id for e in elements if e.text == "Repeated inserted boilerplate."]
    assert ids_first == ids_second
