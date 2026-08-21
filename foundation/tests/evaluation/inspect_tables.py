"""
Deep table and section mapper for Local File Roll-Forward Forensic Audit
"""
import json
from pathlib import Path
import re
import sys

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

REPO_ROOT = Path(__file__).resolve().parents[3]
FOUNDATION_ROOT = REPO_ROOT / "foundation"
sys.path.insert(0, str(REPO_ROOT))
sys.path.insert(0, str(FOUNDATION_ROOT))

import docx
from perception.parser import extract_cell_visible_text
from perception.anchor_builder import build_table_hash
PATH_HIST = REPO_ROOT / "anonymize client/Demo files/Demo files/Compare LF/HMV-24-Final-Local File for FY2023-EN-R0303KPMG.docx"
PATH_TMPL = REPO_ROOT / "anonymize client/Demo files/Demo files/Compare LF/Client-25-Template-Local File for FY20XX-Manufacturer-EN-RddmmKPMG-13062025 (Decree 20-2025).docx"
PATH_GT = REPO_ROOT / "anonymize client/Demo files/Demo files/Compare LF/HMV-26-Final-Local File for FY2024-EN-R2901KPMG.docx"

doc_h = docx.Document(str(PATH_HIST))
doc_t = docx.Document(str(PATH_TMPL))
doc_g = docx.Document(str(PATH_GT))

print("="*80)
print("ALL TABLES IN TEMPLATE (16 tables):")
for i, t in enumerate(doc_t.tables):
    h_text = " | ".join([extract_cell_visible_text(c)[0].strip() for c in t.rows[0].cells][:6])
    h_hash = build_table_hash(t)
    print(f"Tmpl Table {i:2d} (hash={h_hash}, rows={len(t.rows):2d}, cols={len(t.columns):2d}): {h_text}")

print("\n"+"="*80)
print("ALL TABLES IN HISTORICAL FY2023 (22 tables):")
for i, t in enumerate(doc_h.tables):
    h_text = " | ".join([extract_cell_visible_text(c)[0].strip() for c in t.rows[0].cells][:6])
    h_hash = build_table_hash(t)
    print(f"Hist Table {i:2d} (hash={h_hash}, rows={len(t.rows):2d}, cols={len(t.columns):2d}): {h_text}")

print("\n"+"="*80)
print("ALL TABLES IN GROUND TRUTH FY2024 (19 tables):")
for i, t in enumerate(doc_g.tables):
    h_text = " | ".join([extract_cell_visible_text(c)[0].strip() for c in t.rows[0].cells][:6])
    h_hash = build_table_hash(t)
    print(f"GT Table {i:2d} (hash={h_hash}, rows={len(t.rows):2d}, cols={len(t.columns):2d}): {h_text}")
