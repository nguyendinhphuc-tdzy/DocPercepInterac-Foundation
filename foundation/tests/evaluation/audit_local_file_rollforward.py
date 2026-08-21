"""
Local File Roll-Forward Forensic Domain & Capability Audit Script
================================================================
Location: foundation/tests/evaluation/audit_local_file_rollforward.py

Performs deep structural, element-level, table-topology, and media analysis
across the 5 real-world fixtures to establish the exact capability baseline
and domain model for the Local File Roll-Forward workflow.
"""
from __future__ import annotations

import hashlib
import json
import os
from pathlib import Path
import re
import sys
from typing import Any, Dict, List, Optional, Tuple
import zipfile

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")

import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.table import Table, _Cell
import openpyxl

# Add repo root and foundation to sys.path
REPO_ROOT = Path(__file__).resolve().parents[3]
FOUNDATION_ROOT = REPO_ROOT / "foundation"
sys.path.insert(0, str(REPO_ROOT))
sys.path.insert(0, str(FOUNDATION_ROOT))

from perception.parser import extract_geometry, extract_media_manifest, extract_worksheet_metadata, extract_cell_visible_text
from perception.anchor_builder import assign_anchors, build_table_hash
from perception.element_classifier import classify_blocks
from perception.models import Element, ElementType, ElementIndex

# File Paths
PATH_HIST = REPO_ROOT / "anonymize client/Demo files/Demo files/Compare LF/HMV-24-Final-Local File for FY2023-EN-R0303KPMG.docx"
PATH_TMPL = REPO_ROOT / "anonymize client/Demo files/Demo files/Compare LF/Client-25-Template-Local File for FY20XX-Manufacturer-EN-RddmmKPMG-13062025 (Decree 20-2025).docx"
PATH_DATA_FARPT = REPO_ROOT / "anonymize client/Demo files/Demo files/FA&RPTS & Appendix I/FA&RPTs/HMV-FA&RPT FY2024.xlsx"
PATH_DATA_APP1 = REPO_ROOT / "anonymize client/Demo files/Demo files/FA&RPTS & Appendix I/Appendix I/HMV-25-Appendix I under D20 for FY2024-Final-W3103.xlsx"
PATH_GT = REPO_ROOT / "anonymize client/Demo files/Demo files/Compare LF/HMV-26-Final-Local File for FY2024-EN-R2901KPMG.docx"


def run_perception_pipeline(file_path: Path) -> ElementIndex:
    """Runs the exact Foundation perception pipeline."""
    fmt = file_path.suffix.lstrip(".").lower()
    blocks = extract_geometry(str(file_path))
    anchors = assign_anchors(blocks, fmt)
    elements = classify_blocks(blocks, fmt, anchors)
    media = extract_media_manifest(str(file_path), fmt)
    worksheets = extract_worksheet_metadata(str(file_path)) if fmt == "xlsx" else []
    return ElementIndex(
        doc_id="audit-" + file_path.stem[:12],
        source_path=str(file_path),
        format=fmt,
        elements=elements,
        media=media,
        worksheets=worksheets,
    )


def audit_docx_tables_deep(doc_path: Path) -> List[Dict[str, Any]]:
    """Deep inspection of DOCX tables, row counts, cell text, and merged cell properties."""
    doc = docx.Document(str(doc_path))
    tables_info = []
    
    for t_idx, table in enumerate(doc.tables):
        t_hash = build_table_hash(table)
        row_count = len(table.rows)
        col_count = len(table.columns) if row_count > 0 else 0
        
        # Header text
        header_cells = [extract_cell_visible_text(c)[0].strip() for c in table.rows[0].cells] if row_count > 0 else []
        header_summary = " | ".join(header_cells[:8])
        
        # Check merged cell properties (gridSpan, vMerge) and cell details
        has_gridspan = False
        has_vmerge = False
        all_row_texts = []
        
        for r_idx, row in enumerate(table.rows):
            row_texts = []
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                if tcPr.find(qn('w:gridSpan')) is not None:
                    has_gridspan = True
                if tcPr.find(qn('w:vMerge')) is not None:
                    has_vmerge = True
                txt = extract_cell_visible_text(cell)[0].strip()
                row_texts.append(txt)
            all_row_texts.append(row_texts)
                
        tables_info.append({
            "table_index": t_idx,
            "table_hash": t_hash,
            "row_count": row_count,
            "col_count": col_count,
            "header_summary": header_summary,
            "has_gridspan": has_gridspan,
            "has_vmerge": has_vmerge,
            "rows": all_row_texts,
        })
        
    return tables_info


def audit_xlsx_workbook_deep(xlsx_path: Path) -> Dict[str, Any]:
    """Deep inspection of XLSX workbook sheets, formulas, cells, merged ranges, and tables."""
    wb = openpyxl.load_workbook(str(xlsx_path), data_only=False)
    sheets_info = {}
    
    for name in wb.sheetnames:
        ws = wb[name]
        populated_count = 0
        formula_count = 0
        cell_data = []
        
        for r_idx, row in enumerate(ws.iter_rows(values_only=False), start=1):
            row_items = []
            for c_idx, cell in enumerate(row, start=1):
                if cell.value is not None:
                    populated_count += 1
                    val_str = str(cell.value)
                    if val_str.startswith("="):
                        formula_count += 1
                    row_items.append((cell.coordinate, val_str))
            if row_items:
                cell_data.append((r_idx, row_items[:10]))
                        
        sheets_info[name] = {
            "max_row": ws.max_row,
            "max_column": ws.max_column,
            "populated_cells": populated_count,
            "formula_cells": formula_count,
            "merged_ranges": [str(m) for m in ws.merged_cells.ranges],
            "hidden": ws.sheet_state != "visible",
            "sample_rows": cell_data[:20],
        }
        
    return {
        "sheet_names": wb.sheetnames,
        "sheets": sheets_info,
    }


def audit_docx_sections_and_placeholders(doc_path: Path) -> List[Dict[str, Any]]:
    """Analyzes DOCX paragraphs, sections, headings, and placeholder tags."""
    doc = docx.Document(str(doc_path))
    sections_list = []
    
    current_section = "PREAMBLE"
    for p_idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        style = p.style.name if p.style else ""
        
        # Check if heading
        is_heading = "Heading" in style or bool(re.match(r"^(SECTION|PART|CHAPTER|\d+(\.\d+)*)\s+", txt, re.IGNORECASE))
        if is_heading and txt:
            current_section = txt[:80]
            
        # Detect placeholders e.g. FY20xx, [insert], <...>, XXX, etc.
        placeholders = []
        for m in re.finditer(r"(\bFY20[xX]{2}\b|\[[^\]]+\]|<[^>]+>|\bXX+\b|\bYY+\b|\bClient\b|\bddmmKPMG\b)", txt):
            placeholders.append(m.group(0))
            
        sections_list.append({
            "paragraph_index": p_idx,
            "section": current_section,
            "style": style,
            "is_heading": is_heading,
            "text": txt[:120],
            "placeholders": placeholders,
        })
        
    return sections_list


def audit_media_assets_deep(doc_path: Path) -> List[Dict[str, Any]]:
    """Inspects embedded images/drawings and relates them to document context."""
    doc = docx.Document(str(doc_path))
    media_items = []
    
    with zipfile.ZipFile(str(doc_path), 'r') as zf:
        media_names = [n for n in zf.namelist() if n.startswith("word/media/")]
        for m in media_names:
            info = zf.getinfo(m)
            data = zf.read(m)
            sha = hashlib.sha256(data).hexdigest()
            ext = Path(m).suffix.lower()
            media_items.append({
                "source_ref": m,
                "size_bytes": info.file_size,
                "sha256": sha,
                "ext": ext,
            })
            
    # Also find paragraphs containing drawings/images
    drawing_paras = []
    for p_idx, p in enumerate(doc.paragraphs):
        xml = p._p.xml
        if "<w:drawing" in xml or "<w:pict" in xml:
            # find surrounding text
            surrounding = p.text.strip() or (doc.paragraphs[p_idx-1].text.strip() if p_idx > 0 else "")
            drawing_paras.append({
                "paragraph_index": p_idx,
                "surrounding_text": surrounding[:100],
            })
            
    return {
        "media_files": media_items,
        "drawing_paragraphs": drawing_paras,
    }


def main():
    print("=" * 75)
    print("LOCAL FILE ROLL-FORWARD FORENSIC AUDIT (REAL FIXTURES)")
    print("=" * 75)
    
    results = {}
    
    paths = {
        "A_hist_2023": PATH_HIST,
        "B_tmpl_2025": PATH_TMPL,
        "C_data_farpt_2024": PATH_DATA_FARPT,
        "D_data_app1_2024": PATH_DATA_APP1,
        "E_gt_2024": PATH_GT,
    }
    
    print("\n>>> 1. PERCEPTION PIPELINE COVERAGE AUDIT...")
    perception_results = {}
    for key, p in paths.items():
        elem_idx = run_perception_pipeline(p)
        counts = {}
        for el in elem_idx.elements:
            t = el.type.value if hasattr(el.type, "value") else str(el.type)
            counts[t] = counts.get(t, 0) + 1
        perception_results[key] = {
            "total_elements": len(elem_idx.elements),
            "by_type": counts,
            "media_count": len(elem_idx.media),
            "worksheets_count": len(elem_idx.worksheets),
        }
        print(f"  [+] {key:18}: {len(elem_idx.elements):5} elements | Media: {len(elem_idx.media):2} | Sheets: {len(elem_idx.worksheets):2}")
        print(f"      Types: {counts}")
        
    results["perception"] = perception_results
    
    print("\n>>> 2. DEEP AUDIT OF DOCX TABLES...")
    tables_hist = audit_docx_tables_deep(PATH_HIST)
    tables_tmpl = audit_docx_tables_deep(PATH_TMPL)
    tables_gt = audit_docx_tables_deep(PATH_GT)
    
    print(f"  [+] Hist (FY23) Tables: {len(tables_hist)}")
    print(f"  [+] Template Tables: {len(tables_tmpl)}")
    print(f"  [+] Ground Truth (FY24) Tables: {len(tables_gt)}")
    
    # Analyze Template Tables vs Hist vs Ground Truth
    table_comparison = []
    for idx, t_tmpl in enumerate(tables_tmpl):
        # Match table by header similarity or hash
        match_hist = None
        for th in tables_hist:
            if th["table_hash"] == t_tmpl["table_hash"] or (th["header_summary"] and th["header_summary"] == t_tmpl["header_summary"]):
                match_hist = th
                break
        if not match_hist and idx < len(tables_hist):
            match_hist = tables_hist[idx]
            
        match_gt = None
        for tg in tables_gt:
            if tg["table_hash"] == t_tmpl["table_hash"] or (tg["header_summary"] and tg["header_summary"] == t_tmpl["header_summary"]):
                match_gt = tg
                break
        if not match_gt and idx < len(tables_gt):
            match_gt = tables_gt[idx]
            
        hist_rows = match_hist["row_count"] if match_hist else None
        tmpl_rows = t_tmpl["row_count"]
        gt_rows = match_gt["row_count"] if match_gt else None
        
        diff_str = "SAME"
        if hist_rows is not None and gt_rows is not None:
            if gt_rows > hist_rows:
                diff_str = f"GROWTH: {hist_rows} -> {gt_rows} (+{gt_rows-hist_rows})"
            elif gt_rows < hist_rows:
                diff_str = f"SHRINK: {hist_rows} -> {gt_rows} (-{hist_rows-gt_rows})"
            else:
                diff_str = f"EQUAL: {hist_rows}"
                
        table_comparison.append({
            "template_index": idx,
            "template_hash": t_tmpl["table_hash"],
            "header": t_tmpl["header_summary"][:75],
            "tmpl_rows": tmpl_rows,
            "hist_rows": hist_rows,
            "gt_rows": gt_rows,
            "growth_status": diff_str,
            "has_gridspan": t_tmpl["has_gridspan"],
            "has_vmerge": t_tmpl["has_vmerge"],
            "tmpl_sample_header": t_tmpl["rows"][0] if t_tmpl["rows"] else [],
            "gt_sample_header": match_gt["rows"][0] if match_gt and match_gt["rows"] else [],
        })
        
    print(f"\n>>> 3. TABLE ROW GROWTH BREAKDOWN ({len(table_comparison)} Template Tables):")
    for tc in table_comparison:
        print(f"  Table {tc['template_index']:2}: {tc['growth_status']:20} | Tmpl:{tc['tmpl_rows']:2} | Hist:{str(tc['hist_rows']):2} | GT:{str(tc['gt_rows']):2} | {tc['header']}")
        
    results["table_comparison"] = table_comparison
    
    print("\n>>> 4. DEEP AUDIT OF XLSX DATA SOURCES (FA&RPT + APPENDIX I)...")
    farpt_info = audit_xlsx_workbook_deep(PATH_DATA_FARPT)
    app1_info = audit_xlsx_workbook_deep(PATH_DATA_APP1)
    
    print(f"  [+] FA&RPT 2024 ({len(farpt_info['sheet_names'])} sheets):")
    for sname in farpt_info["sheet_names"]:
        sd = farpt_info["sheets"][sname]
        print(f"      - {sname:20}: {sd['populated_cells']:4} cells ({sd['formula_cells']:3} formulas) | Grid: {sd['max_row']}x{sd['max_column']}")
        
    print(f"  [+] Appendix I 2024 ({len(app1_info['sheet_names'])} sheets):")
    for sname in app1_info["sheet_names"]:
        sd = app1_info["sheets"][sname]
        print(f"      - {sname:35}: {sd['populated_cells']:5} cells ({sd['formula_cells']:3} formulas) | Grid: {sd['max_row']}x{sd['max_column']}")
        
    results["xlsx_farpt"] = farpt_info
    results["xlsx_app1"] = app1_info
    
    print("\n>>> 5. DEEP AUDIT OF TEMPLATE PLACEHOLDERS & SECTIONS...")
    tmpl_sections = audit_docx_sections_and_placeholders(PATH_TMPL)
    placeholders_found = {}
    for s in tmpl_sections:
        for ph in s["placeholders"]:
            placeholders_found[ph] = placeholders_found.get(ph, 0) + 1
            
    print(f"  [+] Total Paragraphs in Template: {len(tmpl_sections)}")
    print(f"  [+] Distinct Placeholders Detected: {len(placeholders_found)}")
    print("      Top Placeholders:")
    for ph, count in sorted(placeholders_found.items(), key=lambda x: x[1], reverse=True)[:15]:
        print(f"        '{ph}': {count} occurrences")
        
    results["template_sections"] = tmpl_sections
    results["placeholders_summary"] = placeholders_found
    
    print("\n>>> 6. DEEP AUDIT OF MEDIA, IMAGES & FIGURES...")
    media_hist = audit_media_assets_deep(PATH_HIST)
    media_tmpl = audit_media_assets_deep(PATH_TMPL)
    media_gt = audit_media_assets_deep(PATH_GT)
    
    print(f"  [+] Hist (FY23) Media: {len(media_hist['media_files'])} assets in zip, {len(media_hist['drawing_paragraphs'])} drawing paragraphs")
    print(f"  [+] Template Media   : {len(media_tmpl['media_files'])} assets in zip, {len(media_tmpl['drawing_paragraphs'])} drawing paragraphs")
    print(f"  [+] Ground Truth FY24: {len(media_gt['media_files'])} assets in zip, {len(media_gt['drawing_paragraphs'])} drawing paragraphs")
    
    results["media_audit"] = {
        "hist": media_hist,
        "tmpl": media_tmpl,
        "gt": media_gt,
    }
    
    # Save full audit dump to scratch
    scratch_dir = REPO_ROOT / ".scratch"
    scratch_dir.mkdir(exist_ok=True)
    out_json = scratch_dir / "local_file_rollforward_audit_complete.json"
    out_json.write_text(json.dumps(results, indent=2, default=str), encoding="utf-8")
    print(f"\n[+] Full audit JSON written to: {out_json}")
    print("=" * 75)


if __name__ == "__main__":
    main()
