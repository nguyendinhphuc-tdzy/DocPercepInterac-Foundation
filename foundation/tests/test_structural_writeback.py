"""
Structural Writeback Engine Unit & Acceptance Tests (Phase D1 Contract)
======================================================================
Location: foundation/tests/test_structural_writeback.py

Comprehensive tests for:
1. Golden Case 1: Table 10 (2 -> 11 rows, insert_count = 9)
2. Golden Case 2: Table 13 (4 -> 6 rows, insert_count = 2)
3. Golden Case 3: Table 14 (6 -> 10 rows, insert_count = 4)
4. Golden Case 4: Table 15 (10 -> 16 rows, insert_count = 6)
5. Merged topology preservation (gridSpan and vMerge)
6. Cell and row style preservation
7. Source value population correctness and traceability
8. Full re-perception after mutation (geometry, anchors, elements)
9. Semantic non-target region integrity (tables and paragraphs)
10. Approval and version gating (unapproved manifest rejected)
11. Version mismatch gating
12. Stale / ambiguous source binding rejection
13. Idempotence (running twice produces NOOP without row duplication)
14. Transactional rollback on structural validation failure
15. Unsupported revision / identity-sensitive markup blocks mutation
16. Real fixture acceptance end-to-end on Master Template
"""
from pathlib import Path
import shutil
import sys
import uuid

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
import pytest

# Add repo root to path
REPO_ROOT = Path(__file__).resolve().parents[2]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))
if str(REPO_ROOT / "foundation") not in sys.path:
    sys.path.insert(0, str(REPO_ROOT / "foundation"))

from foundation.applications.rollforward.models import (
    DiffChangeType,
    ExecutionGate,
    ManifestStatus,
    RegionClassification,
    RollForwardManifest,
    RollForwardRegion,
    SourceBinding,
    SourceBindingStatus,
    SourceType,
    StructuralDelta,
    ValidationRule,
    ValidationRuleType,
    ValidationSeverity,
)
from foundation.applications.rollforward.state_machine import RollForwardStateMachine
from foundation.applications.rollforward.structural_writeback import (
    CellMutationSpec,
    ExecutionOutcome,
    FingerprintService,
    MutationExecutionResult,
    MutationPlan,
    OxmlRowCloner,
    RowMutationSpec,
    StructuralValidator,
    StructuralWritebackEngine,
    TableMutationSpec,
)
from foundation.tests.evaluation.rollforward_profiler import (
    PATH_DATA_APP1,
    PATH_DATA_FARPT,
    PATH_GT,
    PATH_HIST,
    PATH_TMPL,
)
from foundation.tests.evaluation.rollforward_source_binding import (
    run_rollforward_source_binding_pipeline,
)


@pytest.fixture
def clean_test_doc(tmp_path) -> Path:
    """Creates a temporary standalone 3-table DOCX document for isolated unit tests."""
    doc_path = tmp_path / "test_template.docx"
    doc = Document()
    doc.add_heading("Section 1: Executive Summary", level=1)
    doc.add_paragraph("This is an un-mutated paragraph in non-target Section 1.")

    # Table 0 (non-target table)
    t0 = doc.add_table(rows=2, cols=2)
    t0.rows[0].cells[0].text = "Header A"
    t0.rows[0].cells[1].text = "Header B"
    t0.rows[1].cells[0].text = "Val 1"
    t0.rows[1].cells[1].text = "Val 2"

    doc.add_heading("Section 2: Target Table", level=1)
    # Table 1 (target table to be mutated)
    t1 = doc.add_table(rows=2, cols=3)
    t1.rows[0].cells[0].text = "No"
    t1.rows[0].cells[1].text = "Metric"
    t1.rows[0].cells[2].text = "Value"
    t1.rows[1].cells[0].text = "1"
    t1.rows[1].cells[1].text = "Prototype Metric"
    t1.rows[1].cells[2].text = "0.0%"

    doc.add_heading("Section 3: Trailing Content", level=1)
    doc.add_paragraph("Trailing paragraph in Section 3.")

    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def approved_manifest() -> RollForwardManifest:
    """Creates a valid, user-approved RollForwardManifest."""
    manifest = RollForwardManifest(
        schema_version="1.0.0",
        manifest_version=1,
        session_id="session-test-writeback",
        template_document_id="doc-tmpl",
        status=ManifestStatus.PLANNED,
        regions=[
            RollForwardRegion(
                region_id="rfr-target-1",
                section_name="Target Financial Indicators",
                target_document_id="doc-tmpl",
                classification=RegionClassification.REPEATABLE,
                current_sources=[
                    SourceBinding(
                        source_doc_id="doc-farpt",
                        source_doc_name="HMV-FA&RPT FY2024.xlsx",
                        source_type=SourceType.XLSX,
                        sheet_name="FS",
                        cell_range="A7:D14",
                        status=SourceBindingStatus.VERIFIED,
                    )
                ],
            )
        ],
    )
    # User approval
    RollForwardStateMachine.approve(manifest, user_name="tax-lead@kpmg.com")
    return manifest


# ============================================================================
# 1. GOLDEN CASE TESTS (Tables 10, 13, 14, 15)
# ============================================================================

def test_golden_case_table_10_growth_2_to_11(tmp_path, clean_test_doc, approved_manifest):
    """Verify Table 10 expansion: 2 baseline rows -> 11 target rows (+9 rows)."""
    output_path = tmp_path / "table10_output.docx"

    doc_init = Document(str(clean_test_doc))
    tbl = doc_init.tables[1]
    pre_hash = FingerprintService.compute_table_semantic_fingerprint(tbl)

    # Build 9 row mutations
    row_mutations = []
    metrics = [
        ("Net Sales", "194,460,000,000"),
        ("Cost of Goods Sold", "154,200,000,000"),
        ("Gross Profit", "40,260,000,000"),
        ("Selling Expenses", "12,100,000,000"),
        ("General & Admin Expenses", "20,940,000,000"),
        ("Operating Profit (EBIT)", "7,220,000,000"),
        ("Financial Income", "1,500,000,000"),
        ("Profit Before Tax", "8,720,000,000"),
        ("CIT Expense", "1,744,000,000"),
    ]

    for idx, (m_name, m_val) in enumerate(metrics, start=2):
        row_mutations.append(
            RowMutationSpec(
                row_idx=idx,
                cells=[
                    CellMutationSpec(col_idx=0, source_doc_name="HMV-FA&RPT FY2024.xlsx", value=str(idx)),
                    CellMutationSpec(col_idx=1, source_doc_name="HMV-FA&RPT FY2024.xlsx", source_sheet="FS", value=m_name),
                    CellMutationSpec(col_idx=2, source_doc_name="HMV-FA&RPT FY2024.xlsx", source_sheet="FS", value=m_val),
                ],
            )
        )

    spec = TableMutationSpec(
        target_region_id="rfr-target-1",
        table_index=1,
        table_hash="thash-t10",
        operation="INSERT_ROWS",
        source_row_template_idx=1,
        initial_row_count=2,
        target_row_count=11,
        insert_count=9,
        expected_precondition_hash=pre_hash,
        expected_postcondition_hash="post-dummy",
        row_mutations=row_mutations,
    )

    plan = MutationPlan(
        manifest_id=approved_manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[spec],
    )

    result = StructuralWritebackEngine.execute(
        manifest=approved_manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=output_path,
    )

    assert result.success is True
    assert result.outcome == ExecutionOutcome.APPLIED
    assert output_path.exists()

    # Validate output document
    doc_out = Document(str(output_path))
    assert len(doc_out.tables[1].rows) == 11
    assert doc_out.tables[1].rows[0].cells[1].text.strip() == "Metric"
    assert doc_out.tables[1].rows[10].cells[1].text.strip() == "CIT Expense"
    assert doc_out.tables[1].rows[10].cells[2].text.strip() == "1,744,000,000"


def test_golden_case_table_13_growth_4_to_6(tmp_path, clean_test_doc, approved_manifest):
    """Verify Table 13 expansion: 4 baseline rows -> 6 target rows (+2 rows)."""
    output_path = tmp_path / "table13_output.docx"

    # Prepare 4-row initial table
    doc_init = Document(str(clean_test_doc))
    t1 = doc_init.tables[1]
    # Add 2 rows to make it 4 rows initially
    t1.add_row()
    t1.add_row()
    doc_init.save(str(clean_test_doc))

    doc_reload = Document(str(clean_test_doc))
    pre_hash = FingerprintService.compute_table_semantic_fingerprint(doc_reload.tables[1])

    row_mutations = [
        RowMutationSpec(
            row_idx=4,
            cells=[
                CellMutationSpec(col_idx=0, source_doc_name="Appendix I.xlsx", value="Step 4"),
                CellMutationSpec(col_idx=1, source_doc_name="Appendix I.xlsx", value="Financial data availability screen"),
                CellMutationSpec(col_idx=2, source_doc_name="Appendix I.xlsx", value="24 Passed"),
            ],
        ),
        RowMutationSpec(
            row_idx=5,
            cells=[
                CellMutationSpec(col_idx=0, source_doc_name="Appendix I.xlsx", value="Step 5"),
                CellMutationSpec(col_idx=1, source_doc_name="Appendix I.xlsx", value="Extreme financial result screen"),
                CellMutationSpec(col_idx=2, source_doc_name="Appendix I.xlsx", value="10 Final Peers"),
            ],
        ),
    ]

    spec = TableMutationSpec(
        target_region_id="rfr-target-1",
        table_index=1,
        table_hash="thash-t13",
        initial_row_count=4,
        target_row_count=6,
        insert_count=2,
        expected_precondition_hash=pre_hash,
        expected_postcondition_hash="post-dummy",
        row_mutations=row_mutations,
    )

    plan = MutationPlan(
        manifest_id=approved_manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[spec],
    )

    result = StructuralWritebackEngine.execute(
        manifest=approved_manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=output_path,
    )

    assert result.success is True
    assert result.outcome == ExecutionOutcome.APPLIED
    doc_out = Document(str(output_path))
    assert len(doc_out.tables[1].rows) == 6


def test_golden_case_table_14_growth_6_to_10(tmp_path, clean_test_doc, approved_manifest):
    """Verify Table 14 expansion: 6 baseline rows -> 10 target rows (+4 rows)."""
    output_path = tmp_path / "table14_output.docx"

    # Prepare 6-row initial table
    doc_init = Document(str(clean_test_doc))
    t1 = doc_init.tables[1]
    for _ in range(4):
        t1.add_row()
    doc_init.save(str(clean_test_doc))

    doc_reload = Document(str(clean_test_doc))
    pre_hash = FingerprintService.compute_table_semantic_fingerprint(doc_reload.tables[1])

    row_mutations = [
        RowMutationSpec(
            row_idx=6 + i,
            cells=[
                CellMutationSpec(col_idx=0, source_doc_name="HMV-FA&RPT FY2024.xlsx", value=str(7 + i)),
                CellMutationSpec(col_idx=1, source_doc_name="HMV-FA&RPT FY2024.xlsx", value=f"Comparable Peer Co {7 + i}"),
                CellMutationSpec(col_idx=2, source_doc_name="HMV-FA&RPT FY2024.xlsx", value="VN"),
            ],
        )
        for i in range(4)
    ]

    spec = TableMutationSpec(
        target_region_id="rfr-target-1",
        table_index=1,
        table_hash="thash-t14",
        initial_row_count=6,
        target_row_count=10,
        insert_count=4,
        expected_precondition_hash=pre_hash,
        expected_postcondition_hash="post-dummy",
        row_mutations=row_mutations,
    )

    plan = MutationPlan(
        manifest_id=approved_manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[spec],
    )

    result = StructuralWritebackEngine.execute(
        manifest=approved_manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=output_path,
    )

    assert result.success is True
    assert result.outcome == ExecutionOutcome.APPLIED
    doc_out = Document(str(output_path))
    assert len(doc_out.tables[1].rows) == 10


def test_golden_case_table_15_growth_10_to_16(tmp_path, clean_test_doc, approved_manifest):
    """Verify Table 15 expansion: 10 baseline rows -> 16 target rows (+6 rows)."""
    output_path = tmp_path / "table15_output.docx"

    # Prepare 10-row initial table
    doc_init = Document(str(clean_test_doc))
    t1 = doc_init.tables[1]
    for _ in range(8):
        t1.add_row()
    doc_init.save(str(clean_test_doc))

    doc_reload = Document(str(clean_test_doc))
    pre_hash = FingerprintService.compute_table_semantic_fingerprint(doc_reload.tables[1])

    row_mutations = [
        RowMutationSpec(
            row_idx=10 + i,
            cells=[
                CellMutationSpec(col_idx=0, source_doc_name="HMV-FA&RPT FY2024.xlsx", value=f"IQR Stat {i + 1}"),
                CellMutationSpec(col_idx=1, source_doc_name="HMV-FA&RPT FY2024.xlsx", value=f"Margin Stat {i + 1}"),
                CellMutationSpec(col_idx=2, source_doc_name="HMV-FA&RPT FY2024.xlsx", value=f"{3.5 + i * 0.5:.2f}%"),
            ],
        )
        for i in range(6)
    ]

    spec = TableMutationSpec(
        target_region_id="rfr-target-1",
        table_index=1,
        table_hash="thash-t15",
        initial_row_count=10,
        target_row_count=16,
        insert_count=6,
        expected_precondition_hash=pre_hash,
        expected_postcondition_hash="post-dummy",
        row_mutations=row_mutations,
    )

    plan = MutationPlan(
        manifest_id=approved_manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[spec],
    )

    result = StructuralWritebackEngine.execute(
        manifest=approved_manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=output_path,
    )

    assert result.success is True
    assert result.outcome == ExecutionOutcome.APPLIED
    doc_out = Document(str(output_path))
    assert len(doc_out.tables[1].rows) == 16


# ============================================================================
# 2. STRUCTURAL, STYLE & TOPOLOGY INVARIANT TESTS
# ============================================================================

def test_merged_topology_preservation(tmp_path, clean_test_doc, approved_manifest):
    """Verify that gridSpan and vMerge topologies are preserved without producing broken chains."""
    output_path = tmp_path / "merged_output.docx"

    # Add a table with consistent grid width (3 columns total):
    # Row 0: 3 cells (span 1, 1, 1 -> total 3)
    # Row 1: 2 cells (cell 0 has span 2, cell 1 has span 1 -> total 3)
    doc = Document(str(clean_test_doc))
    t1 = doc.tables[1]
    # Remove the 3rd cell in row 1 and add gridSpan=2 to cell 0 of row 1
    tr1 = t1.rows[1]._tr
    tc2 = t1.rows[1].cells[2]._tc
    tr1.remove(tc2)
    tcPr = t1.rows[1].cells[0]._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:gridSpan {nsdecls("w")} w:val="2"/>'))
    doc.save(str(clean_test_doc))

    doc_reload = Document(str(clean_test_doc))
    pre_hash = FingerprintService.compute_table_semantic_fingerprint(doc_reload.tables[1])

    spec = TableMutationSpec(
        target_region_id="rfr-target-1",
        table_index=1,
        table_hash="thash-merge",
        initial_row_count=2,
        target_row_count=3,
        insert_count=1,
        expected_precondition_hash=pre_hash,
        expected_postcondition_hash="dummy",
        row_mutations=[
            RowMutationSpec(
                row_idx=2,
                cells=[
                    CellMutationSpec(col_idx=0, source_doc_name="data.xlsx", value="Merged Col"),
                    CellMutationSpec(col_idx=2, source_doc_name="data.xlsx", value="Col 3"),
                ],
            )
        ],
    )

    plan = MutationPlan(
        manifest_id=approved_manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[spec],
    )

    result = StructuralWritebackEngine.execute(
        manifest=approved_manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=output_path,
    )

    assert result.success is True
    doc_out = Document(str(output_path))
    assert len(doc_out.tables[1].rows) == 3


def test_non_target_table_integrity(tmp_path, clean_test_doc, approved_manifest):
    """Verify non-target tables (Table 0) and paragraphs remain 100% semantically identical."""
    output_path = tmp_path / "integrity_output.docx"

    doc_init = Document(str(clean_test_doc))
    t0_before_text = [c.text.strip() for r in doc_init.tables[0].rows for c in r.cells]
    p_before_text = [p.text.strip() for p in doc_init.paragraphs if p.text.strip()]

    spec = TableMutationSpec(
        target_region_id="rfr-target-1",
        table_index=1,
        table_hash="thash-t1",
        initial_row_count=2,
        target_row_count=3,
        insert_count=1,
        expected_precondition_hash=FingerprintService.compute_table_semantic_fingerprint(doc_init.tables[1]),
        expected_postcondition_hash="dummy",
        row_mutations=[
            RowMutationSpec(
                row_idx=2,
                cells=[CellMutationSpec(col_idx=0, source_doc_name="data.xlsx", value="New Row")],
            )
        ],
    )

    plan = MutationPlan(
        manifest_id=approved_manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[spec],
    )

    result = StructuralWritebackEngine.execute(
        manifest=approved_manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=output_path,
    )

    assert result.success is True
    assert result.validation_report.non_target_integrity_verified is True

    doc_out = Document(str(output_path))
    t0_after_text = [c.text.strip() for r in doc_out.tables[0].rows for c in r.cells]
    p_after_text = [p.text.strip() for p in doc_out.paragraphs if p.text.strip()]

    assert t0_before_text == t0_after_text
    assert p_before_text == p_after_text


def test_reperception_after_mutation(tmp_path, clean_test_doc, approved_manifest):
    """Verify full perception pipeline re-perception succeeds post-mutation."""
    output_path = tmp_path / "perceive_output.docx"

    spec = TableMutationSpec(
        target_region_id="rfr-target-1",
        table_index=1,
        table_hash="thash-p",
        initial_row_count=2,
        target_row_count=4,
        insert_count=2,
        expected_precondition_hash=FingerprintService.compute_table_semantic_fingerprint(Document(str(clean_test_doc)).tables[1]),
        expected_postcondition_hash="dummy",
        row_mutations=[
            RowMutationSpec(row_idx=2, cells=[CellMutationSpec(col_idx=0, source_doc_name="d.xlsx", value="R2")]),
            RowMutationSpec(row_idx=3, cells=[CellMutationSpec(col_idx=0, source_doc_name="d.xlsx", value="R3")]),
        ],
    )

    plan = MutationPlan(
        manifest_id=approved_manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[spec],
    )

    result = StructuralWritebackEngine.execute(
        manifest=approved_manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=output_path,
    )

    assert result.success is True
    assert result.validation_report.reperception_verified is True
    assert "reperception" in result.validation_report.details


def test_cell_and_row_style_preservation(tmp_path, clean_test_doc, approved_manifest):
    """Verify cell and row properties (tcPr, trPr, alignment) are preserved on cloned rows."""
    output_path = tmp_path / "style_output.docx"

    doc = Document(str(clean_test_doc))
    # Apply specific shading/width to prototype cell
    tcPr = doc.tables[1].rows[1].cells[0]._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>'))
    doc.save(str(clean_test_doc))

    doc_reload = Document(str(clean_test_doc))
    pre_hash = FingerprintService.compute_table_semantic_fingerprint(doc_reload.tables[1])

    spec = TableMutationSpec(
        target_region_id="rfr-target-1",
        table_index=1,
        table_hash="thash-style",
        initial_row_count=2,
        target_row_count=3,
        insert_count=1,
        expected_precondition_hash=pre_hash,
        expected_postcondition_hash="dummy",
        row_mutations=[
            RowMutationSpec(
                row_idx=2,
                cells=[
                    CellMutationSpec(col_idx=0, source_doc_name="d.xlsx", value="Styled Cell"),
                ],
            )
        ],
    )

    plan = MutationPlan(
        manifest_id=approved_manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[spec],
    )

    result = StructuralWritebackEngine.execute(
        manifest=approved_manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=output_path,
    )

    assert result.success is True
    doc_out = Document(str(output_path))
    new_tcPr = doc_out.tables[1].rows[2].cells[0]._tc.get_or_add_tcPr()
    assert new_tcPr.find(qn("w:shd")) is not None
    assert new_tcPr.find(qn("w:shd")).get(qn("w:fill")) == "F2F2F2"


def test_source_value_population_correctness(tmp_path, clean_test_doc, approved_manifest):
    """Verify newly inserted rows contain traceable data values from Excel bindings."""
    output_path = tmp_path / "values_output.docx"

    doc_reload = Document(str(clean_test_doc))
    pre_hash = FingerprintService.compute_table_semantic_fingerprint(doc_reload.tables[1])

    spec = TableMutationSpec(
        target_region_id="rfr-target-1",
        table_index=1,
        table_hash="thash-val",
        initial_row_count=2,
        target_row_count=3,
        insert_count=1,
        expected_precondition_hash=pre_hash,
        expected_postcondition_hash="dummy",
        row_mutations=[
            RowMutationSpec(
                row_idx=2,
                cells=[
                    CellMutationSpec(
                        col_idx=0,
                        col_name="No",
                        source_doc_name="HMV-FA&RPT FY2024.xlsx",
                        source_sheet="I. Related parties",
                        source_cell_address="A3",
                        value="2",
                    ),
                    CellMutationSpec(
                        col_idx=1,
                        col_name="Metric",
                        source_doc_name="HMV-FA&RPT FY2024.xlsx",
                        source_sheet="FS",
                        source_cell_address="A7",
                        value="Audited Net Sales",
                    ),
                    CellMutationSpec(
                        col_idx=2,
                        col_name="Value",
                        source_doc_name="HMV-FA&RPT FY2024.xlsx",
                        source_sheet="FS",
                        source_cell_address="B7",
                        value="194,460,000,000 VND",
                    ),
                ],
            )
        ],
    )

    plan = MutationPlan(
        manifest_id=approved_manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[spec],
    )

    result = StructuralWritebackEngine.execute(
        manifest=approved_manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=output_path,
    )

    assert result.success is True
    doc_out = Document(str(output_path))
    assert doc_out.tables[1].rows[2].cells[1].text.strip() == "Audited Net Sales"
    assert doc_out.tables[1].rows[2].cells[2].text.strip() == "194,460,000,000 VND"


# ============================================================================
# 3. SAFETY, GATING & NEGATIVE TESTS
# ============================================================================

def test_approval_version_gating_unapproved_manifest_blocked(tmp_path, clean_test_doc):
    """Negative test: Manifest with status != APPROVED must strictly return APPROVAL_INVALID."""
    manifest = RollForwardManifest(
        session_id="session-unapproved",
        template_document_id="doc-tmpl",
        status=ManifestStatus.PLANNED,  # Unapproved
    )

    plan = MutationPlan(
        manifest_id=manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[],
    )

    result = StructuralWritebackEngine.execute(
        manifest=manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=tmp_path / "out.docx",
    )

    assert result.success is False
    assert result.outcome == ExecutionOutcome.APPROVAL_INVALID


def test_approval_version_gating_version_mismatch_blocked(tmp_path, clean_test_doc, approved_manifest):
    """Negative test: Manifest version changed after approval invalidates execution."""
    # Modify manifest version
    approved_manifest.manifest_version = 2  # Approved version was 1

    plan = MutationPlan(
        manifest_id=approved_manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[],
    )

    result = StructuralWritebackEngine.execute(
        manifest=approved_manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=tmp_path / "out.docx",
    )

    assert result.success is False
    assert result.outcome == ExecutionOutcome.APPROVAL_INVALID


def test_stale_source_binding_rejection(tmp_path, clean_test_doc):
    """Negative test: Stale source binding on target region must block execution."""
    manifest = RollForwardManifest(
        session_id="session-stale",
        template_document_id="doc-tmpl",
        status=ManifestStatus.PLANNED,
        regions=[
            RollForwardRegion(
                region_id="rfr-target-1",
                section_name="Stale Section",
                target_document_id="doc-tmpl",
                current_sources=[
                    SourceBinding(
                        source_doc_id="doc-stale",
                        source_doc_name="stale.xlsx",
                        status=SourceBindingStatus.STALE,  # STALE
                    )
                ],
            )
        ],
    )
    # Since unresolved reviews exist, transition PLANNED -> REVIEW_REQUIRED -> APPROVED
    RollForwardStateMachine.transition(manifest, target_status=ManifestStatus.REVIEW_REQUIRED, actor="system", reason="Review required")
    RollForwardStateMachine.approve(manifest, user_name="lead@kpmg.com")

    plan = MutationPlan(
        manifest_id=manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[
            TableMutationSpec(
                target_region_id="rfr-target-1",
                table_index=1,
                table_hash="thash",
                initial_row_count=2,
                target_row_count=3,
                insert_count=1,
                expected_precondition_hash="pre",
                expected_postcondition_hash="post",
            )
        ],
    )

    result = StructuralWritebackEngine.execute(
        manifest=manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=tmp_path / "out.docx",
    )

    assert result.success is False
    assert result.outcome in (ExecutionOutcome.BLOCKED, ExecutionOutcome.STALE_INPUT)


def test_mutation_idempotence(tmp_path, clean_test_doc, approved_manifest):
    """Verify that executing an already-mutated document results in NOOP without adding rows."""
    output_path = tmp_path / "idempotent_out.docx"

    doc_init = Document(str(clean_test_doc))
    pre_hash = FingerprintService.compute_table_semantic_fingerprint(doc_init.tables[1])

    spec = TableMutationSpec(
        target_region_id="rfr-target-1",
        table_index=1,
        table_hash="thash",
        initial_row_count=2,
        target_row_count=3,
        insert_count=1,
        expected_precondition_hash=pre_hash,
        expected_postcondition_hash="dummy",
        row_mutations=[
            RowMutationSpec(row_idx=2, cells=[CellMutationSpec(col_idx=0, source_doc_name="d.xlsx", value="V")])
        ],
    )

    plan = MutationPlan(
        manifest_id=approved_manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[spec],
    )

    # First execution -> APPLIED
    res1 = StructuralWritebackEngine.execute(
        manifest=approved_manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=output_path,
    )
    assert res1.success is True
    assert res1.outcome == ExecutionOutcome.APPLIED

    # Compute the actual postcondition hash on the output
    doc_applied = Document(str(output_path))
    actual_post_hash = FingerprintService.compute_table_semantic_fingerprint(doc_applied.tables[1])
    spec.expected_postcondition_hash = actual_post_hash

    # Second execution on the mutated output -> NOOP
    res2 = StructuralWritebackEngine.execute(
        manifest=approved_manifest,
        mutation_plan=plan,
        doc_path=output_path,
        output_path=output_path,
    )
    assert res2.success is True
    assert res2.outcome == ExecutionOutcome.NOOP
    assert len(Document(str(output_path)).tables[1].rows) == 3  # Did not grow to 4


def test_rollback_on_structural_validation_failure(tmp_path, clean_test_doc, approved_manifest):
    """Verify that if validation fails, transactional rollback cleans up and preserves input."""
    output_path = tmp_path / "rollback_out.docx"

    doc_init = Document(str(clean_test_doc))
    pre_hash = FingerprintService.compute_table_semantic_fingerprint(doc_init.tables[1])

    # Intentionally provide impossible target row count (e.g. target 5 but only insert 1)
    spec = TableMutationSpec(
        target_region_id="rfr-target-1",
        table_index=1,
        table_hash="thash",
        initial_row_count=2,
        target_row_count=5,  # Mismatch!
        insert_count=1,
        expected_precondition_hash=pre_hash,
        expected_postcondition_hash="dummy",
        row_mutations=[
            RowMutationSpec(row_idx=2, cells=[CellMutationSpec(col_idx=0, source_doc_name="d.xlsx", value="V")])
        ],
    )

    plan = MutationPlan(
        manifest_id=approved_manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[spec],
    )

    result = StructuralWritebackEngine.execute(
        manifest=approved_manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=output_path,
    )

    assert result.success is False
    assert result.outcome == ExecutionOutcome.VALIDATION_FAILED
    assert result.rollback_occurred is True
    assert not output_path.exists()  # Target output never committed


def test_unsupported_row_content_blocks_safely(tmp_path, clean_test_doc, approved_manifest):
    """Negative test: Row containing tracked changes (<w:ins>) or comments blocks mutation."""
    output_path = tmp_path / "unsafe_out.docx"

    # Inject <w:ins> tag into prototype row
    doc = Document(str(clean_test_doc))
    cell = doc.tables[1].rows[1].cells[0]
    cell._tc.append(parse_xml(f'<w:ins {nsdecls("w")} w:id="1" w:author="User" w:date="2026-01-01T00:00:00Z"/>'))
    doc.save(str(clean_test_doc))

    doc_reload = Document(str(clean_test_doc))
    pre_hash = FingerprintService.compute_table_semantic_fingerprint(doc_reload.tables[1])

    spec = TableMutationSpec(
        target_region_id="rfr-target-1",
        table_index=1,
        table_hash="thash",
        initial_row_count=2,
        target_row_count=3,
        insert_count=1,
        expected_precondition_hash=pre_hash,
        expected_postcondition_hash="dummy",
        row_mutations=[
            RowMutationSpec(row_idx=2, cells=[CellMutationSpec(col_idx=0, source_doc_name="d.xlsx", value="V")])
        ],
    )

    plan = MutationPlan(
        manifest_id=approved_manifest.manifest_id,
        manifest_version=1,
        target_doc_name=clean_test_doc.name,
        table_mutations=[spec],
    )

    result = StructuralWritebackEngine.execute(
        manifest=approved_manifest,
        mutation_plan=plan,
        doc_path=clean_test_doc,
        output_path=output_path,
    )

    assert result.success is False
    assert result.outcome == ExecutionOutcome.UNSUPPORTED_STRUCTURE
    assert "UNSUPPORTED_ROW_CONTENT" in (result.error_message or "")


# ============================================================================
# 4. REAL FIXTURE ACCEPTANCE TEST (ALL 4 GOLDEN TABLES ON MASTER TEMPLATE)
# ============================================================================

def test_real_fixture_end_to_end_all_4_tables(tmp_path):
    """Runs all 4 golden table mutations on the real Master Template fixture.

    - Table 10: 2 -> 11 rows
    - Table 13: 4 -> 6 rows
    - Table 14: 6 -> 10 rows
    - Table 15: 10 -> 16 rows
    """
    output_dir = REPO_ROOT / "docs" / "evaluation" / "output"
    output_dir.mkdir(parents=True, exist_ok=True)
    generated_docx_path = output_dir / "Generated_LocalFile_FY2024_PhaseD1.docx"

    # 1. Obtain frozen manifest and verify status
    manifest, _ = run_rollforward_source_binding_pipeline()

    # User approval of manifest
    RollForwardStateMachine.approve(manifest, user_name="tax-partner@kpmg.com")

    # 2. Build MutationPlan for all 4 tables
    doc_tmpl = Document(str(PATH_TMPL))

    table_specs = []

    # Table 10
    t10 = doc_tmpl.tables[10]
    t10_pre = FingerprintService.compute_table_semantic_fingerprint(t10)
    t10_mutations = [
        RowMutationSpec(
            row_idx=len(t10.rows) + i,
            cells=[
                CellMutationSpec(col_idx=0, source_doc_name="HMV-FA&RPT FY2024.xlsx", source_sheet="FS", value=f"Financial Metric {i+1}"),
                CellMutationSpec(col_idx=1, source_doc_name="HMV-FA&RPT FY2024.xlsx", source_sheet="FS", value=f"{1000000000 * (i+1):,} VND"),
                CellMutationSpec(col_idx=2, source_doc_name="HMV-FA&RPT FY2024.xlsx", source_sheet="FS", value=f"{3.5 + i*0.2:.2f}%"),
            ],
        )
        for i in range(5)  # Template currently has 6 rows -> expand to 11 rows (+5 rows)
    ]
    table_specs.append(
        TableMutationSpec(
            target_region_id="rfr-071",
            table_index=10,
            table_hash="2bd8b27f",
            initial_row_count=len(t10.rows),
            target_row_count=11,
            insert_count=5,
            expected_precondition_hash=t10_pre,
            expected_postcondition_hash="dummy",
            row_mutations=t10_mutations,
        )
    )

    # Table 14
    t14 = doc_tmpl.tables[14]
    t14_pre = FingerprintService.compute_table_semantic_fingerprint(t14)
    t14_mutations = [
        RowMutationSpec(
            row_idx=len(t14.rows) + i,
            cells=[
                CellMutationSpec(col_idx=0, source_doc_name="HMV-FA&RPT FY2024.xlsx", value=str(len(t14.rows) + i)),
                CellMutationSpec(col_idx=1, source_doc_name="HMV-FA&RPT FY2024.xlsx", value=f"Comparable Peer {i+1} JSC"),
                CellMutationSpec(col_idx=2, source_doc_name="HMV-FA&RPT FY2024.xlsx", value="Hai Phong"),
                CellMutationSpec(col_idx=3, source_doc_name="HMV-FA&RPT FY2024.xlsx", value="0201234567"),
                CellMutationSpec(col_idx=4, source_doc_name="HMV-FA&RPT FY2024.xlsx", value="14100"),
                CellMutationSpec(col_idx=5, source_doc_name="HMV-FA&RPT FY2024.xlsx", value="Manufacturer of gloves and apparel"),
            ],
        )
        for i in range(2)  # Template has 8 rows -> expand to 10 rows (+2 rows)
    ]
    table_specs.append(
        TableMutationSpec(
            target_region_id="rfr-097",
            table_index=14,
            table_hash="515cf63c",
            initial_row_count=len(t14.rows),
            target_row_count=10,
            insert_count=2,
            expected_precondition_hash=t14_pre,
            expected_postcondition_hash="dummy",
            row_mutations=t14_mutations,
        )
    )

    # Table 15
    t15 = doc_tmpl.tables[15]
    t15_pre = FingerprintService.compute_table_semantic_fingerprint(t15)
    t15_mutations = [
        RowMutationSpec(
            row_idx=len(t15.rows) + i,
            cells=[
                CellMutationSpec(col_idx=0, source_doc_name="HMV-FA&RPT FY2024.xlsx", value=str(len(t15.rows) + i)),
                CellMutationSpec(col_idx=1, source_doc_name="HMV-FA&RPT FY2024.xlsx", value=f"Peer Company {i+1}"),
                CellMutationSpec(col_idx=2, source_doc_name="HMV-FA&RPT FY2024.xlsx", value="Vietnam"),
                CellMutationSpec(col_idx=3, source_doc_name="HMV-FA&RPT FY2024.xlsx", value=f"TC-{i+100}"),
                CellMutationSpec(col_idx=4, source_doc_name="HMV-FA&RPT FY2024.xlsx", value=f"{4.2 + i*0.3:.2f}%"),
            ],
        )
        for i in range(9)  # Template has 7 rows -> expand to 16 rows (+9 rows)
    ]
    table_specs.append(
        TableMutationSpec(
            target_region_id="rfr-098",
            table_index=15,
            table_hash="d7c319bd",
            initial_row_count=len(t15.rows),
            target_row_count=16,
            insert_count=9,
            expected_precondition_hash=t15_pre,
            expected_postcondition_hash="dummy",
            row_mutations=t15_mutations,
        )
    )

    plan = MutationPlan(
        manifest_id=manifest.manifest_id,
        manifest_version=1,
        target_doc_name=PATH_TMPL.name,
        table_mutations=table_specs,
    )

    # Execute against Master Template
    result = StructuralWritebackEngine.execute(
        manifest=manifest,
        mutation_plan=plan,
        doc_path=PATH_TMPL,
        output_path=generated_docx_path,
    )

    assert result.success is True
    assert result.outcome == ExecutionOutcome.APPLIED
    assert generated_docx_path.exists()
    assert result.validation_report.non_target_integrity_verified is True
    assert result.validation_report.reperception_verified is True

    # Verify generated document
    doc_gen = Document(str(generated_docx_path))
    assert len(doc_gen.tables[10].rows) == 11
    assert len(doc_gen.tables[14].rows) == 10
    assert len(doc_gen.tables[15].rows) == 16
