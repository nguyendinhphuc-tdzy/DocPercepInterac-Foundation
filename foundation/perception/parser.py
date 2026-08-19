"""Parse capability: Geometry Layer (deterministic, no AI).

Docling has been removed (decision confirmed 2026-08-11, per
Foundation_Build_Plan_v4.md mục 0 điểm 4 + STATUS.md) — replaced by
python-docx for DOCX and pdfplumber (+ pdf2image for page rendering) for
PDF. This is the "extract_geometry" step of the v3 pipeline:

    file_intake -> extract_geometry -> group_into_elements -> assign_anchors

group_into_elements lives in element_classifier.py; assign_anchors lives in
anchor_builder.py (perception/anchor_builder.py) — this module only
produces raw geometry blocks, in document order. No classification, no
anchors, no AI.

`extract_geometry()`'s signature (`str -> list[GeometryBlock]`) is a hard
contract: `applications/gpts/mapping_service.py` and `demo_mapper.py` call
it directly and must not be touched by a perception change (STATUS.md
"Quy tắc không được phá vỡ"). Everything this phase adds — images, charts,
drawings, headers/footers, footnotes/endnotes, comments, PDF images/
annotations — is therefore added as MORE entries in that same flat list
(each tagged with an explicit `kind`, see GeometryBlock below), never a
change to what the function returns or accepts. Document-level side data
that isn't a single object in the document (the media manifest, XLSX
worksheet metadata) is deliberately NOT threaded through this list —
see `extract_media_manifest()` / `extract_worksheet_metadata()`, new
functions only `api/routes/documents.py` calls.
"""
from __future__ import annotations

import re
import zipfile
from pathlib import Path
from typing import Any, Optional, TypedDict

import pdfplumber
from docx import Document as DocxDocument

from perception.models import MediaAsset, WorksheetMetadata

EMU_PER_PX = 9525  # 914400 EMU/inch ÷ 96 px/inch — standard OOXML drawing unit conversion


class GeometryBlock(TypedDict):
    text: str
    kind: str  # "paragraph" | "table_cell" | "image" | "chart" | "drawing" | "header" | "footer" | "footnote" | "endnote" | "comment" | "text_line" | "cell" | "pdf_image" | "annotation" | "page_fallback"
    # DOCX fields
    paragraph_index: Optional[int]
    style_id: Optional[str]
    table_index: Optional[int]
    table_hash: Optional[str]
    row_index: Optional[int]
    col_index: Optional[int]
    # PDF fields
    page: Optional[int]
    bbox: Optional[tuple[float, float, float, float]]  # (x0, top, x1, bottom), pt
    page_width: Optional[float]  # pt — for anchor_builder.py's bbox_relative
    page_height: Optional[float]  # pt
    # XLSX fields
    sheet_name: Optional[str]
    cell_address: Optional[str]
    named_range: Optional[str]
    row_label: Optional[str]  # this row's leftmost non-empty cell — for anchor_builder.py's row self-heal
    # Open-ended metadata for the object kinds this phase adds — deliberately
    # a loose dict rather than a dozen new named TypedDict keys, so every
    # pre-existing block-literal above only needed two new keys (`kind`,
    # `extra`) instead of a large mechanical diff across the whole file.
    # Consumers (element_classifier.py) read specific keys by `kind`.
    extra: Optional[dict[str, Any]]


def _base_block(kind: str, text: str = "", **extra: Any) -> GeometryBlock:
    """Every GeometryBlock field defaults to None/empty here so each
    kind-specific builder below only sets the fields that are actually
    meaningful for it — matches the existing style (every field always
    present) without repeating all 14 keys at every call site."""
    return {
        "text": text,
        "kind": kind,
        "paragraph_index": None,
        "style_id": None,
        "table_index": None,
        "table_hash": None,
        "row_index": None,
        "col_index": None,
        "page": None,
        "bbox": None,
        "page_width": None,
        "page_height": None,
        "sheet_name": None,
        "cell_address": None,
        "named_range": None,
        "row_label": None,
        "extra": extra or None,
    }


# --- DOCX ---------------------------------------------------------------

_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_W_INS_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _paragraph_text(para) -> tuple[str, bool]:
    """`para.text`, extended to include text tracked-inserted via `<w:ins>`.

    python-docx's own `Paragraph.text` is `''.join(r.text for r in
    self.runs)`, and `self.runs` comes from `CT_P.r_lst` — an XPath for
    direct `<w:r>` children of `<w:p>` only. A tracked insertion wraps its
    run(s) one level deeper (`<w:p><w:ins><w:r>...</w:r></w:ins></w:p>`),
    so python-docx never sees that text at all — not "sees it but drops
    formatting", genuinely invisible. Confirmed against the real KPMG
    fixture: one paragraph (all runs insertion-wrapped) produced empty
    `.text` and was silently skipped by the `if para.text.strip():` check
    below entirely; two more had their inserted trailing clause missing
    from otherwise-normal `.text`.

    Deleted text (`<w:del>` / `<w:delText>`) is deliberately NOT included
    here: docx-preview's default render (this project doesn't set
    `renderChanges: true`) renders `<w:ins>` content as plain visible text
    and `<w:del>` content as nothing at all — i.e. an "accepted changes"
    view. Perception's text should match what the rendered document
    actually shows, so inserted text belongs in `text` and deleted text
    does not; adding deleted text back in would make this MORE complete
    than the document that's actually displayed to the user.

    Walks `para._p`'s direct children in document order so normal and
    inserted runs interleave correctly (never a separate "append insertions
    at the end" pass), matching a plain `<w:r>` at a given position and
    recursing one level into a `<w:ins>` at that position for its own
    `<w:r>` children. Returns (text, had_insertion) — the flag lets a
    caller record that this text is partly sourced from a tracked change
    without needing a full revision/diff model.
    """
    from docx.text.run import Run

    ins_tag = f"{{{_W_INS_NS}}}ins"
    r_tag = f"{{{_W_INS_NS}}}r"
    parts: list[str] = []
    had_insertion = False
    for child in para._p:
        if child.tag == r_tag:
            parts.append(Run(child, para).text)
        elif child.tag == ins_tag:
            for r in child.findall(r_tag):
                parts.append(Run(r, para).text)
                had_insertion = True
    return "".join(parts), had_insertion
_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_C_NS = "http://schemas.openxmlformats.org/drawingml/2006/chart"
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _docx_drawings_in_paragraph(para) -> list[GeometryBlock]:
    """Every `<w:drawing>` in this paragraph's XML, in run order — covers
    both inline (`wp:inline`) and floating (`wp:anchor`) pictures, since
    both live inside a run inside a paragraph regardless of positioning
    behavior; walking python-docx's already-parsed paragraph list (rather
    than the whole document.xml separately) keeps image position correctly
    interleaved with text without a second document-order pass."""
    blocks: list[GeometryBlock] = []
    for drawing in para._p.findall(f".//{{{_W_NS}}}drawing"):
        is_inline = drawing.find(f"{{{_WP_NS}}}inline") is not None
        # A chart is a graphicFrame whose graphicData uri is the chart
        # namespace — same drawing container shape as a picture, different
        # payload. Check this before assuming "picture".
        chart_ref = drawing.find(f".//{{{_C_NS}}}chart")
        if chart_ref is not None:
            chart_rid = chart_ref.get(f"{{{_R_NS}}}id")
            blocks.append(_base_block(
                "chart", text="",
                relationship_id=chart_rid, inline=is_inline,
            ))
            continue

        blip = drawing.find(f".//{{{_A_NS}}}blip")
        extent = drawing.find(f".//{{{_WP_NS}}}extent")
        docpr = drawing.find(f".//{{{_WP_NS}}}docPr")
        if blip is not None:
            rid = blip.get(f"{{{_R_NS}}}embed")
            width_px = height_px = None
            if extent is not None:
                width_px = round(int(extent.get("cx", "0")) / EMU_PER_PX)
                height_px = round(int(extent.get("cy", "0")) / EMU_PER_PX)
            blocks.append(_base_block(
                "image", text="",
                relationship_id=rid,
                drawing_id=docpr.get("id") if docpr is not None else None,
                width=width_px, height=height_px, inline=is_inline,
            ))
            continue

        # A `<w:drawing>` that's neither a recognized picture nor a chart —
        # a shape, text box, SmartArt, OLE object, etc. Never silently
        # dropped: emit it as a detected-but-unclassified drawing.
        docpr2 = drawing.find(f".//{{{_WP_NS}}}docPr")
        blocks.append(_base_block(
            "drawing", text="",
            drawing_id=docpr2.get("id") if docpr2 is not None else None,
            name=docpr2.get("name") if docpr2 is not None else None, inline=is_inline,
        ))
    return blocks


def _docx_headers_footers(doc) -> list[GeometryBlock]:
    blocks: list[GeometryBlock] = []
    for s_idx, section in enumerate(doc.sections):
        for kind, part in (("header", section.header), ("footer", section.footer)):
            text = "\n".join(p.text for p in part.paragraphs if p.text.strip())
            if not text:
                continue
            blocks.append(_base_block(kind, text=text, section_index=s_idx))
    return blocks


_NOTE_SKIP_IDS = {"0", "-1"}  # separator/continuation-separator placeholders, not real notes


def _docx_notes_and_comments(path: str) -> list[GeometryBlock]:
    """Footnotes/endnotes/comments have no python-docx API — read the raw
    OOXML parts directly (word/footnotes.xml, endnotes.xml, comments.xml),
    which is exactly the "direct OOXML package inspection" this phase
    calls for beyond what python-docx exposes."""
    import xml.etree.ElementTree as ET

    blocks: list[GeometryBlock] = []
    part_specs = [
        ("word/footnotes.xml", "footnote", f"{{{_W_NS}}}footnote"),
        ("word/endnotes.xml", "endnote", f"{{{_W_NS}}}endnote"),
        ("word/comments.xml", "comment", f"{{{_W_NS}}}comment"),
    ]
    try:
        with zipfile.ZipFile(path) as zf:
            names = set(zf.namelist())
            for part_name, kind, tag in part_specs:
                if part_name not in names:
                    continue
                root = ET.fromstring(zf.read(part_name))
                for note in root.findall(tag):
                    note_id = note.get(f"{{{_W_NS}}}id")
                    if note_id in _NOTE_SKIP_IDS:
                        continue
                    text = "".join(node.text or "" for node in note.iter(f"{{{_W_NS}}}t")).strip()
                    if not text:
                        continue
                    author = note.get(f"{{{_W_NS}}}author") if kind == "comment" else None
                    blocks.append(_base_block(kind, text=text, note_id=note_id, author=author))
    except (KeyError, zipfile.BadZipFile, ET.ParseError):
        pass  # malformed/absent part — never let a chrome-content read failure break the whole parse
    return blocks


def parse_docx(path: str) -> list[GeometryBlock]:
    """Deterministic DOCX geometry via python-docx, extended with direct
    OOXML inspection for object types python-docx has no API for (images,
    charts, drawings, footnotes, endnotes, comments — see module docstring
    for why this is additive to the existing paragraph/table-cell blocks
    rather than a second return channel).

    One block per non-empty paragraph (in body order) plus one block per
    table cell — including empty ones, which are real fill-in placeholders,
    not noise — tagged with the same indices AnchorDOCX expects
    (perception/models.py). Images/charts/drawings are interleaved at the
    paragraph position they actually occur at, preserving true reading
    order (phase requirement: never bucket them into a separate list).
    Headers/footers/footnotes/endnotes/comments are genuinely outside body
    reading order (separate OOXML content streams) and are appended after.
    """
    doc = DocxDocument(path)
    blocks: list[GeometryBlock] = []

    for i, para in enumerate(doc.paragraphs):
        text, had_insertion = _paragraph_text(para)
        if text.strip():
            block = _base_block("paragraph", text=text)
            block["paragraph_index"] = i
            block["style_id"] = para.style.style_id if para.style else None
            block["extra"] = {"has_tracked_insertion": True} if had_insertion else None
            blocks.append(block)
        # Images/charts/drawings inside this paragraph — checked regardless
        # of whether the paragraph itself has text, since a picture-only
        # paragraph has empty `para.text`.
        for drawing_block in _docx_drawings_in_paragraph(para):
            drawing_block["paragraph_index"] = i
            blocks.append(drawing_block)

    for t_idx, table in enumerate(doc.tables):
        # Compute Anti-Drift hash for this table
        from perception.anchor_builder import build_table_hash
        t_hash = build_table_hash(table)

        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                # Unlike paragraphs, empty table cells are kept — a blank
                # cell in a financial template is a real placeholder a
                # user (or a mapping rule) is meant to fill in, not noise.
                block = _base_block("table_cell", text=cell.text)
                block["table_index"] = t_idx
                block["table_hash"] = t_hash
                block["row_index"] = r_idx
                block["col_index"] = c_idx
                blocks.append(block)

    blocks.extend(_docx_headers_footers(doc))
    blocks.extend(_docx_notes_and_comments(path))

    return blocks


# --- PDF ------------------------------------------------------------------


def _pdf_images(page, page_num: int) -> list[GeometryBlock]:
    blocks: list[GeometryBlock] = []
    for i, img in enumerate(page.images):
        block = _base_block(
            "pdf_image", text="",
            image_index=i, width=img.get("width"), height=img.get("height"),
        )
        block["page"] = page_num
        block["bbox"] = (img["x0"], img["top"], img["x1"], img["bottom"])
        block["page_width"] = page.width
        block["page_height"] = page.height
        blocks.append(block)
    return blocks


def _pdf_annotations(page, page_num: int) -> list[GeometryBlock]:
    """Hyperlinks first (pdfplumber's well-supported, stable API), then any
    other annotation subtype pdfplumber exposes via the lower-level
    `page.annots` — wrapped defensively since annotation completeness
    varies by PDF producer and pdfplumber version; a page with unparseable
    annotation dicts must not fail the whole document's perception."""
    blocks: list[GeometryBlock] = []
    for link in page.hyperlinks:
        block = _base_block("annotation", text=link.get("uri", ""), subtype="link", uri=link.get("uri"))
        block["page"] = page_num
        block["bbox"] = (link["x0"], link["top"], link["x1"], link["bottom"])
        block["page_width"] = page.width
        block["page_height"] = page.height
        blocks.append(block)

    try:
        annots = page.annots or []
    except Exception:
        annots = []
    hyperlink_bboxes = {(round(b["bbox"][0], 1), round(b["bbox"][1], 1)) for b in blocks}
    for annot in annots:
        subtype = str(annot.get("data", {}).get("Subtype", "unknown")).strip("/").lower()
        if subtype == "link":
            continue  # already captured via page.hyperlinks with a proper uri
        x0, top, x1, bottom = annot.get("x0"), annot.get("top"), annot.get("x1"), annot.get("bottom")
        if None in (x0, top, x1, bottom):
            continue
        if (round(x0, 1), round(top, 1)) in hyperlink_bboxes:
            continue
        block = _base_block("annotation", text=annot.get("contents") or "", subtype=subtype or "unknown")
        block["page"] = page_num
        block["bbox"] = (x0, top, x1, bottom)
        block["page_width"] = page.width
        block["page_height"] = page.height
        blocks.append(block)
    return blocks


def parse_pdf(path: str) -> list[GeometryBlock]:
    """Deterministic PDF geometry via pdfplumber — one block per text line,
    with its bounding box in PDF points (top-left origin, matches AnchorPDF).
    Extended with image and annotation/hyperlink detection (same pdfplumber
    page object, no extra library). A page with neither extractable text
    nor images gets an explicit `page_fallback` block instead of silently
    contributing zero elements — the phase requirement that a scanned page
    still be represented, not treated as empty (PDF.js still renders the
    actual page raster regardless; this only concerns perception).

    Text extraction only covers a real text layer. Scanned/image-only PDFs
    (no text layer) return zero *text* blocks — OCR is a separate, not-yet-
    decided concern (see STATUS.md fixture_report.pdf note).
    """
    blocks: list[GeometryBlock] = []
    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_had_text = False
            for line in page.extract_text_lines():
                text = line["text"]
                if not text.strip():
                    continue
                page_had_text = True
                block = _base_block("text_line", text=text)
                block["page"] = page_num
                block["bbox"] = (line["x0"], line["top"], line["x1"], line["bottom"])
                block["page_width"] = page.width
                block["page_height"] = page.height
                blocks.append(block)

            image_blocks = _pdf_images(page, page_num)
            blocks.extend(image_blocks)
            blocks.extend(_pdf_annotations(page, page_num))

            if not page_had_text and not image_blocks:
                block = _base_block(
                    "page_fallback",
                    text=f"Page {page_num} — no extractable text or embedded images (likely a scanned page).",
                )
                block["page"] = page_num
                block["bbox"] = (0, 0, page.width, page.height)
                block["page_width"] = page.width
                block["page_height"] = page.height
                blocks.append(block)
    return blocks


# --- XLSX -------------------------------------------------------------------


def _xlsx_images(ws, sheet_name: str) -> list[GeometryBlock]:
    from openpyxl.utils import get_column_letter

    blocks: list[GeometryBlock] = []
    for i, img in enumerate(getattr(ws, "_images", []) or []):
        anchor = img.anchor
        from_cell = to_cell = None
        try:
            frm = anchor._from
            from_cell = f"{get_column_letter(frm.col + 1)}{frm.row + 1}"
            to = getattr(anchor, "to", None)
            if to is not None:
                to_cell = f"{get_column_letter(to.col + 1)}{to.row + 1}"
        except AttributeError:
            pass  # a OneCellAnchor/absolute anchor shape openpyxl represents differently — position stays best-effort
        block = _base_block(
            "image", text="",
            drawing_id=f"{sheet_name}:image:{i}",
            from_cell=from_cell, to_cell=to_cell,
            width=getattr(img, "width", None), height=getattr(img, "height", None),
        )
        block["sheet_name"] = sheet_name
        block["cell_address"] = from_cell or "A1"
        blocks.append(block)
    return blocks


def _xlsx_charts(ws, sheet_name: str) -> list[GeometryBlock]:
    from openpyxl.utils import get_column_letter

    blocks: list[GeometryBlock] = []
    for i, chart in enumerate(getattr(ws, "_charts", []) or []):
        anchor = getattr(chart, "anchor", None)
        from_cell = None
        try:
            frm = anchor._from
            from_cell = f"{get_column_letter(frm.col + 1)}{frm.row + 1}"
        except AttributeError:
            pass
        title = None
        try:
            if chart.title and chart.title.tx and chart.title.tx.rich:
                title = "".join(
                    run.t or "" for para in chart.title.tx.rich.p for run in (para.r or [])
                ).strip() or None
        except AttributeError:
            pass
        block = _base_block(
            "chart", text=title or "",
            drawing_id=f"{sheet_name}:chart:{i}",
            from_cell=from_cell, chart_type=type(chart).__name__, chart_title=title,
        )
        block["sheet_name"] = sheet_name
        block["cell_address"] = from_cell or "A1"
        blocks.append(block)
    return blocks


def parse_xlsx(path: str) -> list[GeometryBlock]:
    """Deterministic XLSX geometry via openpyxl.

    Reads all sheets and non-empty cells, plus worksheet images/charts as
    first-class blocks. Formulas are preserved (loaded with
    `data_only=False`) rather than only ever showing the cached computed
    value — a second `data_only=True` load provides the displayed value
    alongside the formula for cells that have one, without losing either.
    """
    import openpyxl
    from openpyxl.utils import range_boundaries

    blocks: list[GeometryBlock] = []
    wb = openpyxl.load_workbook(path, data_only=False, read_only=False)
    wb_values = openpyxl.load_workbook(path, data_only=True, read_only=False)

    # 1. Build a map of cell coordinates to named ranges
    named_ranges_map: dict[str, list[str]] = {}
    if wb.defined_names:
        for dn in wb.defined_names.values():
            try:
                destinations = list(dn.destinations)
                for sheet_name, coord in destinations:
                    # coord could be a range like "A1:C3" or single cell "A1"
                    min_col, min_row, max_col, max_row = range_boundaries(coord)
                    # To be fully exhaustive, map every cell in the range to this name
                    for row in range(min_row, max_row + 1):
                        for col in range(min_col, max_col + 1):
                            cell_address = openpyxl.utils.get_column_letter(col) + str(row)
                            key = f"{sheet_name}!{cell_address}"
                            if key not in named_ranges_map:
                                named_ranges_map[key] = []
                            named_ranges_map[key].append(dn.name)
            except Exception:
                # Skip legacy/corrupt named ranges that openpyxl cannot parse
                pass

    # 2. Extract cells
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        ws_values = wb_values[sheet_name]
        for row, row_values in zip(ws.iter_rows(), ws_values.iter_rows()):
            # Leftmost non-empty cell in the row — the typical financial-
            # statement convention (line-item label in column A/B, values
            # to the right). Shared by every cell in this row so
            # anchor_builder.py can self-heal row insert/delete drift.
            row_label = next(
                (str(c.value).strip() for c in row_values if c.value is not None), None
            )
            for cell, cell_value in zip(row, row_values):
                if cell.value is None:
                    continue
                is_formula = isinstance(cell.value, str) and cell.value.startswith("=")
                # Check for named ranges
                key = f"{sheet_name}!{cell.coordinate}"
                nr_list = named_ranges_map.get(key, [])
                nr_str = ",".join(nr_list) if nr_list else None

                display_text = str(cell_value.value).strip() if cell_value.value is not None else str(cell.value).strip()
                block = _base_block(
                    "cell", text=display_text,
                    formula=cell.value if is_formula else None,
                    number_format=cell.number_format if cell.number_format != "General" else None,
                    hyperlink=cell.hyperlink.target if cell.hyperlink else None,
                    comment=cell.comment.text.strip() if cell.comment and cell.comment.text else None,
                )
                block["sheet_name"] = sheet_name
                block["cell_address"] = cell.coordinate
                block["named_range"] = nr_str
                block["row_label"] = row_label
                blocks.append(block)

        blocks.extend(_xlsx_images(ws, sheet_name))
        blocks.extend(_xlsx_charts(ws, sheet_name))

    wb.close()
    wb_values.close()
    return blocks


def extract_worksheet_metadata(path: str) -> list[WorksheetMetadata]:
    """Sheet-level display facts (merges, hidden rows/cols, freeze panes,
    explicit dimensions) an XLSX renderer needs — deliberately NOT part of
    `extract_geometry()`'s block list (these describe the grid, not a
    perceivable document object). Only called by api/routes/documents.py,
    never by extract_geometry()'s existing callers."""
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=False, read_only=False)
    result: list[WorksheetMetadata] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        row_heights = {
            r: dim.height for r, dim in ws.row_dimensions.items()
            if dim.height is not None and not dim.hidden
        }
        hidden_rows = [r for r, dim in ws.row_dimensions.items() if dim.hidden]
        column_widths = {
            c: dim.width for c, dim in ws.column_dimensions.items()
            if dim.width is not None and not dim.hidden
        }
        hidden_columns = [c for c, dim in ws.column_dimensions.items() if dim.hidden]
        result.append(WorksheetMetadata(
            sheet_name=sheet_name,
            merged_ranges=[str(r) for r in ws.merged_cells.ranges],
            hidden_rows=hidden_rows,
            hidden_columns=hidden_columns,
            row_heights=row_heights,
            column_widths=column_widths,
            freeze_panes=ws.freeze_panes,
        ))
    wb.close()
    return result


# --- Media manifest / resolution --------------------------------------------


def extract_media_manifest(path: str, fmt: str) -> list[MediaAsset]:
    """Metadata-only inventory of embedded raster images this document
    contains — never the binary itself (see api/routes/documents.py's media
    endpoint, which resolves a `media_id` from this exact manifest back to
    bytes on demand). Charts have no manifest entry: they're vector-
    described OOXML objects, not a single stored image to serve.
    `media_id` is deterministic and re-derivable from the file alone
    (docx: the relationship id; xlsx: sheet+index), so it never needs to be
    persisted anywhere beyond the Element that references it."""
    if fmt == "docx":
        return _docx_media_manifest(path)
    if fmt == "xlsx":
        return _xlsx_media_manifest(path)
    return []  # PDF pages render their own images directly via pdf.js — no separate manifest needed


def _docx_media_manifest(path: str) -> list[MediaAsset]:
    doc = DocxDocument(path)
    assets: list[MediaAsset] = []
    for rid, rel in doc.part.rels.items():
        if "image" not in rel.reltype:
            continue
        try:
            part = rel.target_part
            assets.append(MediaAsset(
                media_id=rid,
                type="image",
                mime_type=part.content_type,
                source_reference=f"docx-rel:{rid}",
            ))
        except (KeyError, ValueError):
            continue  # an external (non-package) image relationship — nothing to serve
    return assets


def _xlsx_media_manifest(path: str) -> list[MediaAsset]:
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=True, read_only=False)
    assets: list[MediaAsset] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for i, img in enumerate(getattr(ws, "_images", []) or []):
            media_id = f"{sheet_name}:image:{i}"
            fmt_ext = (getattr(img, "format", None) or "png").lower()
            mime = {"png": "image/png", "jpeg": "image/jpeg", "jpg": "image/jpeg", "gif": "image/gif"}.get(fmt_ext, "application/octet-stream")
            assets.append(MediaAsset(
                media_id=media_id, type="image", mime_type=mime,
                width=getattr(img, "width", None), height=getattr(img, "height", None),
                source_reference=f"xlsx-drawing:{media_id}",
            ))
    wb.close()
    return assets


def resolve_media_bytes(path: str, fmt: str, media_id: str) -> Optional[tuple[bytes, str]]:
    """Resolves a `media_id` (as listed in `extract_media_manifest()`'s
    output for this same file) back to raw bytes + MIME type. Only ever
    looks inside this one document's own package — `media_id` is opaque
    and re-validated against a manifest freshly computed from the file
    itself, never a filesystem path the caller could point elsewhere."""
    if fmt == "docx":
        doc = DocxDocument(path)
        rel = doc.part.rels.get(media_id)
        if rel is None or "image" not in rel.reltype:
            return None
        try:
            part = rel.target_part
            return part.blob, part.content_type
        except (KeyError, ValueError):
            return None

    if fmt == "xlsx":
        import openpyxl

        m = re.match(r"^(.+):image:(\d+)$", media_id)
        if not m:
            return None
        sheet_name, idx = m.group(1), int(m.group(2))
        wb = openpyxl.load_workbook(path, data_only=True, read_only=False)
        try:
            if sheet_name not in wb.sheetnames:
                return None
            ws = wb[sheet_name]
            images = getattr(ws, "_images", []) or []
            if idx >= len(images):
                return None
            img = images[idx]
            data = img._data()
            fmt_ext = (getattr(img, "format", None) or "png").lower()
            mime = {"png": "image/png", "jpeg": "image/jpeg", "jpg": "image/jpeg", "gif": "image/gif"}.get(fmt_ext, "application/octet-stream")
            return data, mime
        finally:
            wb.close()

    return None


def render_pdf_pages(path: str, dpi: int = 150):
    """Render PDF pages to images for the Input Viewer (frontend Pane 1).

    Requires Poppler (pdftoppm/pdftocairo) on PATH — this is an OS-level
    binary, not a pip package. As of the last check it is NOT installed on
    this dev machine (raises pdf2image.exceptions.PDFInfoNotInstalledError).
    See STATUS.md before relying on this in a demo.
    """
    from pdf2image import convert_from_path

    return convert_from_path(path, dpi=dpi)


def extract_geometry(path: str) -> list[GeometryBlock]:
    """Entry point of the Geometry Layer — dispatches by file extension."""
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        return parse_docx(path)
    if ext == ".xlsx":
        return parse_xlsx(path)
    if ext == ".pdf":
        return parse_pdf(path)
    raise ValueError(f"Unsupported format for geometry extraction: {ext}")
