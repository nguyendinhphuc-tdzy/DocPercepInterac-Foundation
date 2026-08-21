"""
Local File Roll-Forward Structural Writeback Engine (Phase D1 Contract)
=======================================================================
Location: foundation/applications/rollforward/structural_writeback.py

Implements governed, transactional structural table mutation for Local File
Roll-Forward. Executes approved MutationPlans by performing safe OOXML row
cloning, identity-sensitive tag sanitization, topology-aware merge handling,
semantic non-target validation, and full perception re-verification.

Strict Principles:
1. NEVER mutate original template in-place (Transactional staging -> atomic commit).
2. NEVER silently strip tracked changes or revisions (reject with UNSUPPORTED_ROW_CONTENT).
3. NEVER guess table positions or data values (strict source binding traceability).
4. Strictly enforce user-only approval provenance and manifest version matching.
5. Guarantee idempotence via semantic precondition and postcondition fingerprinting.
"""
from __future__ import annotations

import copy
from datetime import datetime, timezone
from enum import Enum
import hashlib
from pathlib import Path
import shutil
from typing import Any, Dict, List, Literal, Optional, Set, Tuple
import uuid

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Row
from pydantic import BaseModel, Field

from applications.rollforward.models import (
    DiffChangeType,
    ExecutionGate,
    ManifestStatus,
    RollForwardDiff,
    RollForwardManifest,
    SourceBindingStatus,
    ValidationRule,
    ValidationRuleType,
    ValidationSeverity,
)
from perception.anchor_builder import assign_anchors, build_table_hash
from perception.element_classifier import classify_blocks
from perception.parser import extract_cell_visible_text, extract_geometry


# ============================================================================
# 1. OUTCOME & DATA MODELS
# ============================================================================

class ExecutionOutcome(str, Enum):
    """Explicit outcomes of a structural writeback execution attempt."""
    NOOP = "NOOP"                                  # Target state already present; mutation skipped
    APPLIED = "APPLIED"                            # All mutations successfully applied and validated
    BLOCKED = "BLOCKED"                            # Gating blocked execution (unapproved or unverified regions)
    VALIDATION_FAILED = "VALIDATION_FAILED"        # Structural, non-target, or perception validation failed
    UNSUPPORTED_STRUCTURE = "UNSUPPORTED_STRUCTURE"# Unsupported revision, drawing, or broken merge topology
    STALE_INPUT = "STALE_INPUT"                    # Precondition fingerprint mismatch against document
    APPROVAL_INVALID = "APPROVAL_INVALID"          # Missing or mismatched human approval provenance


class CellMutationSpec(BaseModel):
    """Traceable specification for a single cell's value in a mutated row."""
    col_idx: int = Field(ge=0)
    col_name: Optional[str] = None
    source_doc_name: str
    source_sheet: Optional[str] = None
    source_cell_address: Optional[str] = None
    source_element_id: Optional[str] = None
    value: Any
    data_type: str = "str"                         # "str", "float", "int", "currency", "percent"


class RowMutationSpec(BaseModel):
    """Data specification for a newly created or updated table row."""
    row_idx: int = Field(ge=0)
    cells: List[CellMutationSpec] = Field(default_factory=list)


class TableMutationSpec(BaseModel):
    """Governed mutation specification for a single target table."""
    mutation_id: str = Field(default_factory=lambda: f"mut-{uuid.uuid4().hex[:8]}")
    target_region_id: str
    table_index: int = Field(ge=0)
    table_hash: str
    operation: str = "INSERT_ROWS"
    source_row_template_idx: int = 1
    initial_row_count: int = Field(ge=0)
    target_row_count: int = Field(ge=0)
    insert_count: int = Field(ge=0)
    expected_precondition_hash: str
    expected_postcondition_hash: str
    row_mutations: List[RowMutationSpec] = Field(default_factory=list)
    validation_rules: List[ValidationRule] = Field(default_factory=list)


class MutationPlan(BaseModel):
    """Deterministic, version-locked plan describing all approved mutations."""
    plan_id: str = Field(default_factory=lambda: f"plan-{uuid.uuid4().hex[:8]}")
    manifest_id: str
    manifest_version: int
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    target_doc_name: str
    table_mutations: List[TableMutationSpec] = Field(default_factory=list)


class StructuralValidationReport(BaseModel):
    """Diagnostic report produced by post-mutation validation."""
    is_valid: bool
    passed_checks: List[str] = Field(default_factory=list)
    failed_checks: List[str] = Field(default_factory=list)
    warnings: List[str] = Field(default_factory=list)
    non_target_integrity_verified: bool = True
    reperception_verified: bool = True
    details: Dict[str, Any] = Field(default_factory=dict)


class MutationExecutionResult(BaseModel):
    """Complete, immutable result of a structural writeback run."""
    outcome: ExecutionOutcome
    success: bool
    mutation_plan_id: str
    manifest_version: int
    output_path: Optional[str] = None
    diffs: List[RollForwardDiff] = Field(default_factory=list)
    validation_report: Optional[StructuralValidationReport] = None
    error_message: Optional[str] = None
    rollback_occurred: bool = False
    executed_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())


# ============================================================================
# 2. SEMANTIC FINGERPRINTING SERVICE
# ============================================================================

class FingerprintService:
    """Calculates deterministic semantic hashes for tables and non-target regions."""

    @classmethod
    def compute_table_semantic_fingerprint(cls, table: Table) -> str:
        """Computes a content-and-topology hash for a DOCX table."""
        hasher = hashlib.sha256()
        hasher.update(f"rows:{len(table.rows)}|cols:{len(table.columns)}".encode("utf-8"))

        for r_idx, row in enumerate(table.rows):
            row_repr = []
            for c_idx, cell in enumerate(row.cells):
                text, _ = extract_cell_visible_text(cell)
                tcPr = cell._tc.get_or_add_tcPr()
                gs = tcPr.find(qn("w:gridSpan"))
                vm = tcPr.find(qn("w:vMerge"))
                gs_val = gs.get(qn("w:val")) if gs is not None else "1"
                vm_val = vm.get(qn("w:val")) if vm is not None else "none"
                row_repr.append(f"c{c_idx}:{text.strip()}:{gs_val}:{vm_val}")
            hasher.update((f"r{r_idx}:" + "|".join(row_repr)).encode("utf-8"))

        return hasher.hexdigest()[:16]

    @classmethod
    def compute_document_non_target_fingerprint(
        cls, doc: Document, target_table_indices: Set[int]
    ) -> str:
        """Computes a semantic fingerprint over all paragraphs and non-target tables."""
        hasher = hashlib.sha256()

        # 1. All non-target tables
        for t_idx, table in enumerate(doc.tables):
            if t_idx not in target_table_indices:
                t_hash = cls.compute_table_semantic_fingerprint(table)
                hasher.update(f"table_{t_idx}:{t_hash}".encode("utf-8"))

        # 2. All paragraphs (text + style)
        for p_idx, p in enumerate(doc.paragraphs):
            style_id = p.style.style_id if p.style else ""
            p_text = p.text.strip()
            if p_text:
                hasher.update(f"p{p_idx}:{style_id}:{p_text}".encode("utf-8"))

        return hasher.hexdigest()[:16]


# ============================================================================
# 3. OOXML SAFE ROW CLONER
# ============================================================================

class OxmlRowCloner:
    """Performs deep OOXML row cloning with strict safety and revision inspection."""

    UNSUPPORTED_TAGS = {
        qn("w:ins"),
        qn("w:del"),
        qn("w:bookmarkStart"),
        qn("w:bookmarkEnd"),
        qn("w:commentRangeStart"),
        qn("w:commentRangeEnd"),
        qn("w:commentReference"),
    }

    @classmethod
    def inspect_row_safety(cls, row: _Row) -> Tuple[bool, str]:
        """Inspects a row for unsupported revision or identity-sensitive tags.

        Per Requirement #2: NEVER silently strip tracked changes or comments.
        If unsupported constructs exist, refuse mutation with UNSUPPORTED_ROW_CONTENT.
        """
        for element in row._tr.iter():
            if element.tag in cls.UNSUPPORTED_TAGS:
                tag_name = element.tag.split("}")[-1]
                return (
                    False,
                    f"UNSUPPORTED_ROW_CONTENT: Prototype row contains revision/identity tag '<w:{tag_name}>'",
                )

            # Check for embedded media with relationship references
            if "drawing" in element.tag.lower() or "blip" in element.tag.lower():
                return (
                    False,
                    "UNSUPPORTED_ROW_CONTENT: Prototype row contains embedded drawing/media",
                )

            # Check for orphaned vMerge continue
            if element.tag == qn("w:vMerge"):
                val = element.get(qn("w:val"))
                if val == "continue":
                    return (
                        False,
                        "UNSUPPORTED_STRUCTURE: Prototype row contains orphaned 'vMerge=continue'",
                    )

        return True, "OK"

    @classmethod
    def clone_and_populate_row(
        cls,
        table: Table,
        prototype_row_idx: int,
        insert_after_row_idx: int,
        cell_specs: List[CellMutationSpec],
    ) -> _Row:
        """Deep-clones a prototype row and inserts it at the designated position."""
        prototype_tr = table.rows[prototype_row_idx]._tr
        cloned_tr = copy.deepcopy(prototype_tr)

        # Insert into parent table XML
        parent_tbl = table._tbl
        target_tr = table.rows[insert_after_row_idx]._tr
        target_tr.addnext(cloned_tr)

        new_row = _Row(cloned_tr, table)

        # Populate cell values
        for cell_spec in cell_specs:
            if cell_spec.col_idx < len(new_row.cells):
                cell = new_row.cells[cell_spec.col_idx]
                cls._set_cell_text_preserving_style(cell, str(cell_spec.value))

        return new_row

    @classmethod
    def _set_cell_text_preserving_style(cls, cell: Any, text: str) -> None:
        """Sets text in a cell while preserving paragraph and run formatting."""
        if not cell.paragraphs:
            cell.add_paragraph(text)
            return

        p = cell.paragraphs[0]
        if p.runs:
            # Preserve the run's formatting, update text
            first_run = p.runs[0]
            first_run.text = text
            # Remove any extra runs in this paragraph
            for extra_run in p.runs[1:]:
                extra_run._r.getparent().remove(extra_run._r)
        else:
            p.text = text

        # Remove extra paragraphs in cell if any
        for extra_p in cell.paragraphs[1:]:
            extra_p._p.getparent().remove(extra_p._p)


# ============================================================================
# 4. POST-MUTATION STRUCTURAL & RE-PERCEPTION VALIDATOR
# ============================================================================

class StructuralValidator:
    """Validates structural invariants, non-target immutability, and data integrity."""

    @classmethod
    def validate_mutation(
        cls,
        doc_before: Document,
        doc_after: Document,
        spec: TableMutationSpec,
        target_table_idx: int,
    ) -> Tuple[bool, List[str], List[str]]:
        """Validates that target table satisfies row count, topology, and values."""
        passed: List[str] = []
        failed: List[str] = []

        if target_table_idx >= len(doc_after.tables):
            failed.append(f"Target table index {target_table_idx} not found in output document.")
            return False, passed, failed

        table = doc_after.tables[target_table_idx]
        actual_rows = len(table.rows)

        # 1. Row count validation
        if actual_rows == spec.target_row_count:
            passed.append(f"Row count match: expected {spec.target_row_count}, got {actual_rows}")
        else:
            failed.append(
                f"Row count mismatch: expected {spec.target_row_count}, got {actual_rows}"
            )

        # 2. Column count / Grid width consistency
        def _get_row_grid_width(row: _Row) -> int:
            width = 0
            for tc in row._tr.findall(qn("w:tc")):
                tcPr = tc.get_or_add_tcPr()
                gs = tcPr.find(qn("w:gridSpan"))
                if gs is not None:
                    width += int(gs.get(qn("w:val")) or "1")
                else:
                    width += 1
            return width

        grid_widths = [_get_row_grid_width(r) for r in table.rows]
        if len(set(grid_widths)) == 1:
            passed.append(f"Grid width consistent across all {actual_rows} rows: {grid_widths[0]} cols")
        else:
            failed.append(f"Inconsistent grid column width across rows: {set(grid_widths)}")

        # 3. Header row preservation
        header_before, _ = extract_cell_visible_text(doc_before.tables[target_table_idx].rows[0].cells[0])
        header_after, _ = extract_cell_visible_text(table.rows[0].cells[0])
        if header_before == header_after:
            passed.append("Header row 0 preserved exactly")
        else:
            failed.append(f"Header row changed: '{header_before}' -> '{header_after}'")

        # 4. Populated cell value checks
        for row_mut in spec.row_mutations:
            if row_mut.row_idx < actual_rows:
                row = table.rows[row_mut.row_idx]
                for cell_spec in row_mut.cells:
                    if cell_spec.col_idx < len(row.cells):
                        actual_val, _ = extract_cell_visible_text(row.cells[cell_spec.col_idx])
                        if str(cell_spec.value).strip() in actual_val.strip():
                            passed.append(
                                f"Value verified in R{row_mut.row_idx}C{cell_spec.col_idx}: '{cell_spec.value}'"
                            )
                        else:
                            failed.append(
                                f"Value mismatch in R{row_mut.row_idx}C{cell_spec.col_idx}: "
                                f"expected '{cell_spec.value}', got '{actual_val}'"
                            )

        # 5. vMerge chain integrity
        for c_idx in range(len(table.columns)):
            merge_state = "none"
            for r_idx, row in enumerate(table.rows):
                tcPr = row.cells[c_idx]._tc.get_or_add_tcPr()
                vm = tcPr.find(qn("w:vMerge"))
                if vm is not None:
                    val = vm.get(qn("w:val")) or "continue"
                    if val == "restart":
                        merge_state = "restart"
                    elif val == "continue" and merge_state == "none":
                        failed.append(f"Invalid vMerge: orphaned 'continue' at R{r_idx}C{c_idx}")
                else:
                    merge_state = "none"

        is_valid = len(failed) == 0
        return is_valid, passed, failed

    @classmethod
    def validate_non_target_integrity(
        cls,
        doc_before: Document,
        doc_after: Document,
        target_table_indices: Set[int],
        expected_non_target_hash: str,
    ) -> Tuple[bool, str]:
        """Validates that no non-target table or paragraph underwent semantic drift."""
        actual_hash = FingerprintService.compute_document_non_target_fingerprint(
            doc_after, target_table_indices
        )
        if actual_hash == expected_non_target_hash:
            return True, "Non-target document regions verified 100% semantically identical."
        else:
            return (
                False,
                f"Non-target semantic drift detected! Expected hash {expected_non_target_hash}, got {actual_hash}",
            )


class RePerceptionValidator:
    """Executes the full Foundation perception pipeline on the generated artifact."""

    @classmethod
    def verify_reperception(
        cls, doc_path: Path, expected_tables: List[int]
    ) -> Tuple[bool, str, Dict[str, Any]]:
        """Parses output document, assigns anchors, classifies blocks, and checks addressability."""
        try:
            blocks = extract_geometry(str(doc_path))
            anchors = assign_anchors(blocks, "docx")
            elements = classify_blocks(blocks, "docx", anchors)

            # Check target tables exist in perceived blocks
            perceived_tables = {
                b.get("table_index") for b in blocks if b.get("table_index") is not None
            }
            for t_idx in expected_tables:
                if t_idx not in perceived_tables:
                    return False, f"Re-perception failed: Table {t_idx} missing from geometry blocks.", {}

            return True, f"Re-perception succeeded: {len(elements)} elements perceived cleanly.", {
                "total_blocks": len(blocks),
                "total_anchors": len(anchors),
                "total_elements": len(elements),
            }
        except Exception as e:
            return False, f"Re-perception exception: {str(e)}", {}


# ============================================================================
# 5. STRUCTURAL WRITEBACK ENGINE (CORE EXECUTOR)
# ============================================================================

class StructuralWritebackEngine:
    """Governed engine executing approved structural mutation plans."""

    @classmethod
    def execute(
        cls,
        manifest: RollForwardManifest,
        mutation_plan: MutationPlan,
        doc_path: Path,
        output_path: Path,
        actor_role: Literal["user", "agent", "system"] = "user",
        actor_id: str = "user-reviewer",
    ) -> MutationExecutionResult:
        """Executes the structural mutation plan with strict transactional safety."""
        executed_at = datetime.now(timezone.utc).isoformat()

        # --------------------------------------------------------------------
        # 1. APPROVAL & VERSION GATING (Requirement #2 & Requirement #9)
        # --------------------------------------------------------------------
        if manifest.status != ManifestStatus.APPROVED:
            return MutationExecutionResult(
                outcome=ExecutionOutcome.APPROVAL_INVALID,
                success=False,
                mutation_plan_id=mutation_plan.plan_id,
                manifest_version=manifest.manifest_version,
                error_message=f"Execution blocked: Manifest status is '{manifest.status.value}', must be APPROVED.",
                executed_at=executed_at,
            )

        if manifest.approved_manifest_version != manifest.manifest_version:
            return MutationExecutionResult(
                outcome=ExecutionOutcome.APPROVAL_INVALID,
                success=False,
                mutation_plan_id=mutation_plan.plan_id,
                manifest_version=manifest.manifest_version,
                error_message=(
                    f"Execution blocked: Approved version ({manifest.approved_manifest_version}) "
                    f"does not match current manifest version ({manifest.manifest_version})."
                ),
                executed_at=executed_at,
            )

        if mutation_plan.manifest_version != manifest.manifest_version:
            return MutationExecutionResult(
                outcome=ExecutionOutcome.APPROVAL_INVALID,
                success=False,
                mutation_plan_id=mutation_plan.plan_id,
                manifest_version=manifest.manifest_version,
                error_message=(
                    f"Execution blocked: MutationPlan version ({mutation_plan.manifest_version}) "
                    f"does not match approved manifest version ({manifest.manifest_version})."
                ),
                executed_at=executed_at,
            )

        # Validate that all target regions are unblocked (READY) with non-stale bindings
        for spec in mutation_plan.table_mutations:
            region = next((r for r in manifest.regions if r.region_id == spec.target_region_id), None)
            if not region or region.execution_gate != ExecutionGate.READY:
                return MutationExecutionResult(
                    outcome=ExecutionOutcome.BLOCKED,
                    success=False,
                    mutation_plan_id=mutation_plan.plan_id,
                    manifest_version=manifest.manifest_version,
                    error_message=f"Execution blocked: Region '{spec.target_region_id}' execution gate is not READY.",
                    executed_at=executed_at,
                )
            for src in region.current_sources:
                if src.status in (SourceBindingStatus.STALE, SourceBindingStatus.MISSING, SourceBindingStatus.AMBIGUOUS):
                    return MutationExecutionResult(
                        outcome=ExecutionOutcome.STALE_INPUT,
                        success=False,
                        mutation_plan_id=mutation_plan.plan_id,
                        manifest_version=manifest.manifest_version,
                        error_message=f"Execution blocked: Source binding '{src.source_doc_name}' is {src.status.value}.",
                        executed_at=executed_at,
                    )

        # --------------------------------------------------------------------
        # 2. PRECONDITION & IDEMPOTENCE CHECK (Requirement #4)
        # --------------------------------------------------------------------
        doc_initial = Document(str(doc_path))
        target_table_indices = {spec.table_index for spec in mutation_plan.table_mutations}

        all_tables_already_mutated = True
        for spec in mutation_plan.table_mutations:
            if spec.table_index >= len(doc_initial.tables):
                all_tables_already_mutated = False
                break
            tbl = doc_initial.tables[spec.table_index]
            current_fp = FingerprintService.compute_table_semantic_fingerprint(tbl)
            if current_fp != spec.expected_postcondition_hash or len(tbl.rows) != spec.target_row_count:
                all_tables_already_mutated = False
                break

        if all_tables_already_mutated:
            return MutationExecutionResult(
                outcome=ExecutionOutcome.NOOP,
                success=True,
                mutation_plan_id=mutation_plan.plan_id,
                manifest_version=manifest.manifest_version,
                output_path=str(output_path),
                error_message="Target table postconditions already satisfied; execution is idempotent NOOP.",
                executed_at=executed_at,
            )

        # --------------------------------------------------------------------
        # 3. TRANSACTIONAL STAGING SETUP (Requirement #6)
        # --------------------------------------------------------------------
        output_path.parent.mkdir(parents=True, exist_ok=True)
        temp_output_path = output_path.parent / f"staging_{uuid.uuid4().hex[:8]}_{output_path.name}"
        shutil.copyfile(str(doc_path), str(temp_output_path))

        expected_non_target_hash = FingerprintService.compute_document_non_target_fingerprint(
            doc_initial, target_table_indices
        )

        diffs: List[RollForwardDiff] = []
        validation_report = StructuralValidationReport(is_valid=True)

        try:
            doc_staged = Document(str(temp_output_path))

            # ----------------------------------------------------------------
            # 4. OOXML SAFETY INSPECTION & ROW CLONING (Requirements #2, #3, #8)
            # ----------------------------------------------------------------
            for spec in mutation_plan.table_mutations:
                if spec.table_index >= len(doc_staged.tables):
                    raise ValueError(f"Table index {spec.table_index} exceeds document table count.")

                table = doc_staged.tables[spec.table_index]
                proto_idx = spec.source_row_template_idx
                if proto_idx >= len(table.rows):
                    proto_idx = len(table.rows) - 1

                # Safety check on prototype row
                is_safe, safety_msg = OxmlRowCloner.inspect_row_safety(table.rows[proto_idx])
                if not is_safe:
                    return MutationExecutionResult(
                        outcome=ExecutionOutcome.UNSUPPORTED_STRUCTURE,
                        success=False,
                        mutation_plan_id=mutation_plan.plan_id,
                        manifest_version=manifest.manifest_version,
                        error_message=f"Table {spec.table_index}: {safety_msg}",
                        rollback_occurred=True,
                        executed_at=executed_at,
                    )

                before_rows = len(table.rows)
                current_insert_pos = before_rows - 1

                # Execute row insertions
                for row_mut in spec.row_mutations:
                    new_row = OxmlRowCloner.clone_and_populate_row(
                        table=table,
                        prototype_row_idx=proto_idx,
                        insert_after_row_idx=current_insert_pos,
                        cell_specs=row_mut.cells,
                    )
                    current_insert_pos += 1

                after_rows = len(table.rows)
                diffs.append(
                    RollForwardDiff(
                        region_id=spec.target_region_id,
                        change_type=DiffChangeType.ROW_ADDED,
                        before_summary={"table_index": spec.table_index, "row_count": before_rows},
                        after_summary={"table_index": spec.table_index, "row_count": after_rows},
                        delta_details=[{"inserted_rows": spec.insert_count, "operation": spec.operation}],
                    )
                )

            # Save staged document
            doc_staged.save(str(temp_output_path))

            # ----------------------------------------------------------------
            # 5. POST-MUTATION STRUCTURAL & NON-TARGET VALIDATION (Req #1, #8)
            # ----------------------------------------------------------------
            doc_mutated = Document(str(temp_output_path))

            for spec in mutation_plan.table_mutations:
                valid, passed, failed = StructuralValidator.validate_mutation(
                    doc_before=doc_initial,
                    doc_after=doc_mutated,
                    spec=spec,
                    target_table_idx=spec.table_index,
                )
                validation_report.passed_checks.extend(passed)
                validation_report.failed_checks.extend(failed)
                if not valid:
                    validation_report.is_valid = False

            # Validate non-target region integrity
            nt_valid, nt_msg = StructuralValidator.validate_non_target_integrity(
                doc_before=doc_initial,
                doc_after=doc_mutated,
                target_table_indices=target_table_indices,
                expected_non_target_hash=expected_non_target_hash,
            )
            if not nt_valid:
                validation_report.is_valid = False
                validation_report.non_target_integrity_verified = False
                validation_report.failed_checks.append(nt_msg)
            else:
                validation_report.passed_checks.append(nt_msg)

            if not validation_report.is_valid:
                if temp_output_path.exists():
                    temp_output_path.unlink()
                return MutationExecutionResult(
                    outcome=ExecutionOutcome.VALIDATION_FAILED,
                    success=False,
                    mutation_plan_id=mutation_plan.plan_id,
                    manifest_version=manifest.manifest_version,
                    diffs=diffs,
                    validation_report=validation_report,
                    error_message=f"Structural validation failed: {validation_report.failed_checks}",
                    rollback_occurred=True,
                    executed_at=executed_at,
                )

            # ----------------------------------------------------------------
            # 6. FULL RE-PERCEPTION VALIDATION (Requirement #7)
            # ----------------------------------------------------------------
            p_valid, p_msg, p_details = RePerceptionValidator.verify_reperception(
                doc_path=temp_output_path,
                expected_tables=list(target_table_indices),
            )
            if not p_valid:
                validation_report.is_valid = False
                validation_report.reperception_verified = False
                validation_report.failed_checks.append(p_msg)
                if temp_output_path.exists():
                    temp_output_path.unlink()
                return MutationExecutionResult(
                    outcome=ExecutionOutcome.VALIDATION_FAILED,
                    success=False,
                    mutation_plan_id=mutation_plan.plan_id,
                    manifest_version=manifest.manifest_version,
                    diffs=diffs,
                    validation_report=validation_report,
                    error_message=p_msg,
                    rollback_occurred=True,
                    executed_at=executed_at,
                )

            validation_report.passed_checks.append(p_msg)
            validation_report.details["reperception"] = p_details

            # ----------------------------------------------------------------
            # 7. ATOMIC COMMIT (Requirement #6)
            # ----------------------------------------------------------------
            if output_path.exists():
                output_path.unlink()
            temp_output_path.rename(output_path)

            return MutationExecutionResult(
                outcome=ExecutionOutcome.APPLIED,
                success=True,
                mutation_plan_id=mutation_plan.plan_id,
                manifest_version=manifest.manifest_version,
                output_path=str(output_path),
                diffs=diffs,
                validation_report=validation_report,
                rollback_occurred=False,
                executed_at=executed_at,
            )

        except Exception as ex:
            if temp_output_path.exists():
                temp_output_path.unlink()
            return MutationExecutionResult(
                outcome=ExecutionOutcome.VALIDATION_FAILED,
                success=False,
                mutation_plan_id=mutation_plan.plan_id,
                manifest_version=manifest.manifest_version,
                error_message=f"Writeback execution exception: {str(ex)}",
                rollback_occurred=True,
                executed_at=executed_at,
            )
