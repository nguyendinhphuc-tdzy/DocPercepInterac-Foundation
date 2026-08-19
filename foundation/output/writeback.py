"""Write-back Engine for patching values directly into DOCX and XLSX."""
from pathlib import Path
from typing import Any, List, Optional, Tuple
import re

from perception.anchor_builder import resolve_docx_anchor, resolve_table_anchor, resolve_xlsx_anchor


def _coerce_xlsx_value(val: Any) -> Any:
    """Coerces incoming edit strings to appropriate Python/Excel types while preserving
    formulas, strings with leading zeros (e.g. tax codes / phone numbers), and booleans."""
    if not isinstance(val, str):
        return val
    if val == "":
        return None
    if val.startswith("="):
        return val
    # Integer
    try:
        if val == "0" or not (val.startswith("0") and len(val) > 1):
            return int(val)
    except ValueError:
        pass
    # Float
    try:
        if not (val.startswith("0") and len(val) > 1 and not val.startswith("0.")):
            return float(val)
    except ValueError:
        pass
    # Boolean
    if val.lower() == "true":
        return True
    if val.lower() == "false":
        return False
    return val


class WritebackEngine:
    def __init__(self):
        pass

    def apply_single_patch(self, path: str, anchor: Any, new_value: Any, output_path: str) -> Optional[str]:
        """Resolves a single perception.models.Anchor (AnchorDOCX/AnchorXLSX
        — as opposed to apply_patches_docx/xlsx's hard-coded string-anchor
        batches from GTPS's DEMO_RULES) against the document at `path`,
        writes `new_value` there, saves to `output_path`. Returns the
        resolver's self-heal message (or None) — used by PATCH
        /api/elements/<id> for one-off edits made live in the UI.
        """
        str_value = str(new_value)

        if anchor.format == "docx":
            from docx import Document
            from docx.table import _Cell

            doc = Document(path)
            resolved, message = resolve_docx_anchor(doc, anchor)
            if isinstance(resolved, _Cell):
                self._safe_replace_docx_cell(resolved, str_value)
            else:
                self._safe_replace_docx_para(resolved, str_value)
            doc.save(output_path)
            return message

        if anchor.format == "xlsx":
            import openpyxl

            wb = openpyxl.load_workbook(path)
            cell, message = resolve_xlsx_anchor(wb, anchor)
            if isinstance(cell.value, str) and cell.value.startswith("=") and not (isinstance(new_value, str) and new_value.startswith("=")):
                raise ValueError(
                    f"Cell {anchor.sheet_name}!{anchor.cell_address} contains a formula ('{cell.value}') and is read-only"
                )
            cell.value = _coerce_xlsx_value(new_value)
            wb.save(output_path)
            return message

        raise ValueError(
            f"Editing is not supported for '{anchor.format}' output — no write-back "
            "path exists for this format (PDF is read-only in this project)."
        )

    def apply_patches_docx(self, path: str, patches: List[Tuple[str, Any]], output_path: str):
        """Applies a list of patches (anchor, new_value) to a DOCX file."""
        from docx import Document
        doc = Document(path)

        for anchor, new_value in patches:
            str_value = str(new_value)
            
            if anchor.startswith("table:"):
                # Use Anti-Drift resolver
                cell, message = resolve_table_anchor(doc, anchor)
                if cell:
                    if message:
                        print(f"  [i] {message}")
                    self._safe_replace_docx_cell(cell, str_value)
                else:
                    print(f"  [!] Failed to find {anchor} in {Path(path).name}: {message}")
            elif anchor.startswith("para:"):
                match = re.match(r"para:(\d+)", anchor)
                if match:
                    p_idx = int(match.group(1))
                    try:
                        para = doc.paragraphs[p_idx]
                        self._safe_replace_docx_para(para, str_value)
                    except IndexError:
                        print(f"  [!] Failed to find {anchor} in {Path(path).name}")

        doc.save(output_path)
        print(f"  [+] Saved patched DOCX: {Path(output_path).name}")

    def _safe_replace_docx_cell(self, cell, new_value: str):
        """Replaces text in a DOCX cell while preserving the first run's formatting."""
        if not cell.paragraphs:
            cell.add_paragraph()
            
        first_para = cell.paragraphs[0]
        self._safe_replace_docx_para(first_para, new_value)
        
        # Clear subsequent paragraphs
        for para in cell.paragraphs[1:]:
            for run in para.runs:
                run.text = ""

    def _safe_replace_docx_para(self, para, new_value: str):
        """Replaces text in a DOCX paragraph while preserving the first run's formatting."""
        if not para.runs:
            para.add_run()
            
        # Set text in the first run
        para.runs[0].text = new_value
        
        # Clear all subsequent runs
        for run in para.runs[1:]:
            run.text = ""

    def apply_patches_xlsx(self, path: str, patches: List[Tuple[str, Any]], output_path: str):
        """Applies a list of patches (anchor, new_value) to an XLSX file."""
        import openpyxl
        wb = openpyxl.load_workbook(path)
        
        for anchor, new_value in patches:
            # Parse SheetName!CellAddress
            if "!" in anchor:
                sheet_name, cell_addr = anchor.split("!", 1)
                try:
                    ws = wb[sheet_name]
                    ws[cell_addr].value = new_value
                except KeyError:
                    print(f"  [!] Failed to find {anchor} in {Path(path).name}")
        
        wb.save(output_path)
        print(f"  [+] Saved patched XLSX: {Path(output_path).name}")
