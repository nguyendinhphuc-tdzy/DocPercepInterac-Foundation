import sys
from pathlib import Path
sys.path.insert(0, '.')
sys.path.insert(0, './foundation')

from docx import Document
from tests.evaluation.rollforward_profiler import PATH_TMPL, PATH_HIST, PATH_GT

for name, p in [("TEMPLATE", PATH_TMPL), ("HISTORICAL FY23", PATH_HIST), ("GROUND TRUTH FY24", PATH_GT)]:
    doc = Document(str(p))
    print(f"\n==================== {name} (Total Tables: {len(doc.tables)}) ====================")
    for idx in [10, 13, 14, 15]:
        if idx < len(doc.tables):
            tbl = doc.tables[idx]
            print(f"\n--- TABLE {idx} (Rows: {len(tbl.rows)}, Cols: {len(tbl.columns)}) ---")
            for r_idx, row in enumerate(tbl.rows):
                row_txt = [c.text.strip().replace('\n', ' ') for c in row.cells]
                # remove adjacent identical cells from gridSpan display
                dedup_row = []
                for c in row_txt:
                    if not dedup_row or c != dedup_row[-1]:
                        dedup_row.append(c)
                print(f"  R{r_idx:2d}: {dedup_row[:6]}")
